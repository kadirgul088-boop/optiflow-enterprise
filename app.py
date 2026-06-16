import os
import json
from datetime import datetime
import smtplib
from email.message import EmailMessage

import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
import streamlit as st

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from supabase import create_client

from modules.scoring import calculate_optiflow_score
from modules.financial import calculate_financial_impact
from modules.maturity import get_maturity_comment
from modules.recommendations import generate_recommendations
from modules.ai_engine import generate_consulting_report
from modules.report_engine import create_enterprise_pdf
from modules.ai_copilot import ask_real_ai_copilot

try:
    from modules.excel_engine import process_uploaded_excel, create_template_excel
except Exception:
    process_uploaded_excel = None
    create_template_excel = None

try:
    from modules.ppt_engine import create_enterprise_pptx
except Exception:
    create_enterprise_pptx = None


st.set_page_config(
    page_title="OptiFlow Consulting",
    page_icon="📊",
    layout="wide"
)


PROJECT_DIR = "client_projects"
EXPORT_DIR = "exports"
DATA_DIR = "data"
BENCHMARK_DB_PATH = os.path.join(DATA_DIR, "benchmark_database.json")

os.makedirs(PROJECT_DIR, exist_ok=True)
os.makedirs(EXPORT_DIR, exist_ok=True)
os.makedirs(DATA_DIR, exist_ok=True)


DEFAULT_BENCHMARK_DB = {
    "Tekstil": {"wait_rate": 18, "oee": 74, "defect_rate": 4, "line_balance_loss": 15},
    "Otomotiv": {"wait_rate": 12, "oee": 82, "defect_rate": 2, "line_balance_loss": 10},
    "Gida": {"wait_rate": 15, "oee": 78, "defect_rate": 3, "line_balance_loss": 12},
    "Lojistik": {"wait_rate": 20, "oee": 70, "defect_rate": 5, "line_balance_loss": 18},
    "Genel Uretim": {"wait_rate": 17, "oee": 75, "defect_rate": 4, "line_balance_loss": 14}
}


def ensure_benchmark_db():
    if not os.path.exists(BENCHMARK_DB_PATH):
        with open(BENCHMARK_DB_PATH, "w", encoding="utf-8") as file:
            json.dump(DEFAULT_BENCHMARK_DB, file, ensure_ascii=False, indent=4)


def load_benchmark_db():
    ensure_benchmark_db()
    with open(BENCHMARK_DB_PATH, "r", encoding="utf-8") as file:
        return json.load(file)


def save_benchmark_db(db):
    with open(BENCHMARK_DB_PATH, "w", encoding="utf-8") as file:
        json.dump(db, file, ensure_ascii=False, indent=4)


def money_fmt(value):
    try:
        return f"{float(value):,.0f} TL".replace(",", ".")
    except Exception:
        return f"{value} TL"


def pct_fmt(value):
    try:
        return f"%{float(value):.1f}"
    except Exception:
        return f"%{value}"


def risk_status(wait_rate, oee, defect_rate, line_balance_loss):
    risk_score = round(
        (
            float(wait_rate)
            + float(line_balance_loss)
            + float(defect_rate)
            + max(0, 100 - float(oee))
        ) / 4,
        1
    )

    if risk_score >= 25:
        return risk_score, "Yüksek Risk", "#dc2626", "#fff1f2"

    if risk_score >= 15:
        return risk_score, "Orta Risk", "#d97706", "#fffbeb"

    return risk_score, "Düşük Risk", "#059669", "#ecfdf5"


def kpi_status(label, value):
    value = float(value)

    if label == "wait":
        if value > 20:
            return "Kritik"
        if value > 12:
            return "Dikkat"
        return "İyi"

    if label == "oee":
        if value >= 80:
            return "Güçlü"
        if value >= 65:
            return "Dikkat"
        return "Kritik"

    if label == "defect":
        if value <= 3:
            return "İyi"
        if value <= 7:
            return "Dikkat"
        return "Kritik"

    if label == "balance":
        if value > 15:
            return "Kritik"
        if value > 8:
            return "Dikkat"
        return "İyi"

    if label == "roi":
        if value >= 150:
            return "Güçlü"
        if value >= 75:
            return "Dikkat"
        return "Zayıf"

    return "Normal"


def benchmark_from_database(company_metrics, sector):
    db = load_benchmark_db()
    sector_ref = db.get(sector, db.get("Genel Uretim", {}))

    return {
        "Firma Bekleme Oranı": company_metrics.get("wait_rate", 0),
        "Sektör Bekleme Oranı": sector_ref.get("wait_rate", 0),
        "Firma OEE": company_metrics.get("oee", 0),
        "Sektör OEE": sector_ref.get("oee", 0),
        "Firma Hata Oranı": company_metrics.get("defect_rate", 0),
        "Sektör Hata Oranı": sector_ref.get("defect_rate", 0),
        "Firma Hat Denge Kaybı": company_metrics.get("line_balance_loss", 0),
        "Sektör Hat Denge Kaybı": sector_ref.get("line_balance_loss", 0)
    }


def save_project_record(company_name, sector, score, maturity, company_metrics, financial_result, recommendations):
    safe_company = company_name.replace(" ", "_").replace("/", "_").replace("\\", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{safe_company}_{timestamp}.json"
    path = os.path.join(PROJECT_DIR, filename)

    data = {
        "company_name": company_name,
        "sector": sector,
        "created_at": datetime.now().isoformat(),
        "score": score,
        "maturity": maturity,
        "company_metrics": company_metrics,
        "financial_result": financial_result,
        "recommendations": recommendations
    }

    with open(path, "w", encoding="utf-8") as file:
        json.dump(data, file, ensure_ascii=False, indent=4)

    return path


def load_client_records():
    files = sorted(
        [file for file in os.listdir(PROJECT_DIR) if file.endswith(".json")],
        reverse=True
    )

    records = []

    for file_name in files:
        try:
            with open(os.path.join(PROJECT_DIR, file_name), "r", encoding="utf-8") as file:
                record = json.load(file)

            record["_file"] = file_name
            records.append(record)
        except Exception:
            pass

    return records


def plot_gauge(title, value, color_mode="score"):
    value = float(value)

    if color_mode == "risk":
        steps = [
            {"range": [0, 15], "color": "#dcfce7"},
            {"range": [15, 25], "color": "#fef3c7"},
            {"range": [25, 100], "color": "#fee2e2"}
        ]
        bar_color = "#dc2626" if value >= 25 else "#d97706" if value >= 15 else "#059669"
    else:
        steps = [
            {"range": [0, 50], "color": "#fee2e2"},
            {"range": [50, 75], "color": "#fef3c7"},
            {"range": [75, 100], "color": "#dcfce7"}
        ]
        bar_color = "#059669" if value >= 75 else "#d97706" if value >= 50 else "#dc2626"

    fig = go.Figure(
        go.Indicator(
            mode="gauge+number",
            value=value,
            title={"text": title, "font": {"size": 20}},
            number={"font": {"size": 34}},
            gauge={
                "axis": {"range": [0, 100]},
                "bar": {"color": bar_color, "thickness": 0.35},
                "steps": steps,
                "borderwidth": 1,
                "bordercolor": "#cbd5e1"
            }
        )
    )

    fig.update_layout(
        height=300,
        margin=dict(l=20, r=20, t=55, b=15),
        paper_bgcolor="white",
        font=dict(color="#0f172a")
    )

    return fig


def plot_benchmark(benchmark_result):
    data = pd.DataFrame({
        "KPI": ["Bekleme", "OEE", "Hata", "Denge Kaybı"],
        "Firma": [
            float(benchmark_result.get("Firma Bekleme Oranı", 0)),
            float(benchmark_result.get("Firma OEE", 0)),
            float(benchmark_result.get("Firma Hata Oranı", 0)),
            float(benchmark_result.get("Firma Hat Denge Kaybı", 0))
        ],
        "Sektör": [
            float(benchmark_result.get("Sektör Bekleme Oranı", 0)),
            float(benchmark_result.get("Sektör OEE", 0)),
            float(benchmark_result.get("Sektör Hata Oranı", 0)),
            float(benchmark_result.get("Sektör Hat Denge Kaybı", 0))
        ]
    })

    long_df = data.melt(id_vars="KPI", var_name="Kategori", value_name="Değer")

    fig = px.bar(
        long_df,
        x="KPI",
        y="Değer",
        color="Kategori",
        barmode="group",
        text="Değer",
        color_discrete_map={
            "Firma": "#2563eb",
            "Sektör": "#94a3b8"
        }
    )

    fig.update_layout(
        title="Benchmark Comparison",
        height=360,
        plot_bgcolor="white",
        paper_bgcolor="white",
        margin=dict(l=20, r=20, t=55, b=30),
        legend_title_text=""
    )

    fig.update_traces(texttemplate="%{text:.1f}", textposition="outside")

    return fig


def plot_financial_waterfall(financial_result):
    total_loss = float(financial_result.get("Toplam Operasyonel Kayıp", 0))
    improvement = float(financial_result.get("İyileştirme Potansiyeli", financial_result.get("Tahmini Yıllık Tasarruf", 0)))
    saving = float(financial_result.get("Tahmini Yıllık Tasarruf", 0))

    fig = go.Figure(
        go.Waterfall(
            name="Financial Impact",
            orientation="v",
            measure=["absolute", "relative", "total"],
            x=["Operasyonel Kayıp", "İyileştirme Etkisi", "Net Tasarruf"],
            y=[total_loss, -improvement, saving],
            connector={"line": {"color": "#64748b"}},
            increasing={"marker": {"color": "#dc2626"}},
            decreasing={"marker": {"color": "#059669"}},
            totals={"marker": {"color": "#2563eb"}}
        )
    )

    fig.update_layout(
        title="Financial Impact Waterfall",
        height=360,
        plot_bgcolor="white",
        paper_bgcolor="white",
        margin=dict(l=20, r=20, t=55, b=30)
    )

    return fig


def plot_risk_heatmap(wait_rate, oee, defect_rate, line_balance_loss):
    risk_matrix = pd.DataFrame(
        [
            [float(wait_rate), float(defect_rate)],
            [float(line_balance_loss), max(0, 100 - float(oee))]
        ],
        index=["Akış / Kapasite", "Denge / OEE"],
        columns=["Bekleme Riski", "Kalite / Etkinlik Riski"]
    )

    fig = px.imshow(
        risk_matrix,
        text_auto=True,
        color_continuous_scale=["#dcfce7", "#fef3c7", "#fee2e2"],
        aspect="auto"
    )

    fig.update_layout(
        title="Enterprise Risk Heatmap",
        height=350,
        margin=dict(l=20, r=20, t=55, b=30),
        paper_bgcolor="white"
    )

    return fig


def render_plotly_dashboard(
    company_name,
    score,
    maturity,
    company_metrics,
    benchmark_result,
    financial_result,
    risk_score,
    risk_level
):
    wait_rate = company_metrics.get("wait_rate", 0)
    oee = company_metrics.get("oee", 0)
    defect_rate = company_metrics.get("defect_rate", 0)
    line_balance_loss = company_metrics.get("line_balance_loss", 0)

    st.markdown("## Executive Enterprise Dashboard")

    col1, col2, col3, col4, col5 = st.columns(5)

    with col1:
        st.metric("OptiFlow Score", f"{score}/100")
    with col2:
        st.metric("Risk Level", risk_level)
    with col3:
        st.metric("Annual Saving", money_fmt(financial_result.get("Tahmini Yıllık Tasarruf", 0)))
    with col4:
        st.metric("ROI", pct_fmt(financial_result.get("ROI (%)", 0)))
    with col5:
        st.metric("Payback", f"{financial_result.get('Geri Dönüş Süresi (Ay)', 0)} Ay")

    gauge_col1, gauge_col2 = st.columns(2)

    with gauge_col1:
        st.plotly_chart(plot_gauge("OptiFlow Score", score, "score"), use_container_width=True)

    with gauge_col2:
        st.plotly_chart(plot_gauge("Enterprise Risk", risk_score, "risk"), use_container_width=True)

    st.markdown(
        f"""
        <div style="padding:20px;border-radius:18px;background:#eff6ff;border-left:6px solid #2563eb;margin-bottom:20px;">
        <b>CEO Insight:</b><br>
        {company_name} için analiz sonucu, operasyonel iyileştirme potansiyelinin yönetim seviyesinde takip edilmesi gerektiğini göstermektedir.
        OptiFlow Score <b>{score}/100</b>, risk seviyesi <b>{risk_level}</b> ve yıllık tasarruf potansiyeli
        <b>{money_fmt(financial_result.get("Tahmini Yıllık Tasarruf", 0))}</b> seviyesindedir.
        </div>
        """,
        unsafe_allow_html=True
    )

    chart_col1, chart_col2 = st.columns(2)

    with chart_col1:
        st.plotly_chart(plot_benchmark(benchmark_result), use_container_width=True)

    with chart_col2:
        st.plotly_chart(plot_financial_waterfall(financial_result), use_container_width=True)

    st.plotly_chart(
        plot_risk_heatmap(wait_rate, oee, defect_rate, line_balance_loss),
        use_container_width=True
    )

    st.markdown("### KPI Health Matrix")

    kpi_rows = [
        ["Bekleme Oranı", pct_fmt(wait_rate), kpi_status("wait", wait_rate)],
        ["OEE", pct_fmt(oee), kpi_status("oee", oee)],
        ["Hata / Fire Oranı", pct_fmt(defect_rate), kpi_status("defect", defect_rate)],
        ["Hat Denge Kaybı", pct_fmt(line_balance_loss), kpi_status("balance", line_balance_loss)],
        ["ROI", pct_fmt(financial_result.get("ROI (%)", 0)), kpi_status("roi", financial_result.get("ROI (%)", 0))]
    ]

    st.dataframe(
        pd.DataFrame(kpi_rows, columns=["KPI", "Değer", "Durum"]),
        use_container_width=True,
        hide_index=True
    )

    st.markdown("### Operasyonel Olgunluk")
    st.info(f"{maturity.get('Seviye', '-')} — {maturity.get('Yorum', '-')}")







# ============================================================
# SAAS LAYER: AUTH + PLANS + SUPABASE
# ============================================================

SUPABASE_URL = st.secrets.get("SUPABASE_URL", "")
SUPABASE_ANON_KEY = st.secrets.get("SUPABASE_ANON_KEY", "")
SUPABASE_SERVICE_KEY = st.secrets.get("SUPABASE_SERVICE_KEY", "")


@st.cache_resource
def get_supabase_client():
    key = SUPABASE_SERVICE_KEY or SUPABASE_ANON_KEY
    if not SUPABASE_URL or not key:
        return None
    return create_client(SUPABASE_URL, key)


PLAN_RULES = {
    "demo": {
        "label": "Demo",
        "analysis_limit": 1,
        "pdf": False,
        "ppt": False,
        "excel": False,
        "ai": False,
        "price": "Free Demo"
    },
    "starter": {
        "label": "Starter",
        "analysis_limit": 10,
        "pdf": True,
        "ppt": False,
        "excel": True,
        "ai": False,
        "price": "$49 / month"
    },
    "professional": {
        "label": "Professional",
        "analysis_limit": 100,
        "pdf": True,
        "ppt": True,
        "excel": True,
        "ai": True,
        "price": "$149 / month"
    },
    "enterprise": {
        "label": "Enterprise",
        "analysis_limit": 999999,
        "pdf": True,
        "ppt": True,
        "excel": True,
        "ai": True,
        "price": "Custom"
    }
}


def get_app_url():
    return "https://optiflow-enterprise-ewxicdsl96crsq2ycortqo.streamlit.app"


def is_logged_in():
    try:
        return bool(st.user.is_logged_in)
    except Exception:
        return False


def current_user_email():
    if st.session_state.get("demo_mode", False):
        return "demo@optiflow.ai"

    try:
        return st.user.email
    except Exception:
        return None


def current_user_name():
    if st.session_state.get("demo_mode", False):
        return "Demo User"

    try:
        return st.user.name
    except Exception:
        return current_user_email() or "User"


def get_user_profile(email):
    supabase = get_supabase_client()

    if not supabase or not email:
        return {
            "email": email or "demo@optiflow.ai",
            "full_name": "Demo User",
            "plan": "demo",
            "plan_status": "active",
            "analysis_limit": 1
        }

    try:
        result = supabase.table("users").select("*").eq("email", email).limit(1).execute()
        if result.data:
            return result.data[0]

        # First login: create demo user record automatically.
        payload = {
            "email": email,
            "full_name": current_user_name(),
            "plan": "demo",
            "plan_status": "active",
            "analysis_limit": 1
        }

        supabase.table("users").insert(payload).execute()
        return payload

    except Exception as exc:
        st.warning(f"Supabase user profile error: {exc}")
        return {
            "email": email,
            "full_name": current_user_name(),
            "plan": "demo",
            "plan_status": "active",
            "analysis_limit": 1
        }


def get_subscription(email, fallback_plan="demo"):
    supabase = get_supabase_client()

    if st.session_state.get("demo_mode", False):
        return {
            "plan": "demo",
            "status": "active"
        }

    if not supabase or not email:
        return {
            "plan": fallback_plan,
            "status": "inactive"
        }

    try:
        result = supabase.table("subscriptions").select("*").eq("user_email", email).limit(1).execute()

        if result.data:
            return result.data[0]

        # First login: create inactive/demo subscription.
        payload = {
            "user_email": email,
            "plan": fallback_plan,
            "status": "inactive"
        }

        supabase.table("subscriptions").insert(payload).execute()
        return payload

    except Exception as exc:
        st.warning(f"Supabase subscription error: {exc}")
        return {
            "plan": fallback_plan,
            "status": "inactive"
        }


def resolve_active_plan(email):
    profile = get_user_profile(email)
    subscription = get_subscription(email, profile.get("plan", "demo"))

    profile_plan = str(profile.get("plan", "demo")).lower()
    sub_plan = str(subscription.get("plan", profile_plan)).lower()
    sub_status = str(subscription.get("status", "inactive")).lower()

    # Paid plans only unlock features when subscription status is active/trialing.
    if sub_plan in ["starter", "professional", "enterprise"] and sub_status in ["active", "trialing"]:
        plan = sub_plan
    else:
        plan = profile_plan if profile_plan in PLAN_RULES else "demo"

    if st.session_state.get("demo_mode", False):
        plan = "demo"

    return plan, PLAN_RULES.get(plan, PLAN_RULES["demo"]), profile, subscription


def count_user_projects(email):
    supabase = get_supabase_client()

    if not supabase or not email:
        return 0

    try:
        result = supabase.table("projects").select("id", count="exact").eq("user_email", email).execute()
        return result.count or 0
    except Exception:
        return 0


def can_create_analysis(email, plan_rules):
    used = count_user_projects(email)
    limit = int(plan_rules.get("analysis_limit", 1))
    return used < limit, used, limit


def save_project_to_supabase(
    user_email,
    company_name,
    sector,
    score,
    risk_level,
    financial_result
):
    supabase = get_supabase_client()

    if not supabase or not user_email:
        return None

    try:
        payload = {
            "user_email": user_email,
            "company_name": company_name,
            "sector": sector,
            "score": float(score),
            "risk_level": risk_level,
            "annual_saving": float(financial_result.get("Tahmini Yıllık Tasarruf", 0))
        }

        result = supabase.table("projects").insert(payload).execute()

        if result.data:
            return result.data[0].get("id")

    except Exception as exc:
        st.warning(f"Supabase project save error: {exc}")

    return None


def save_report_to_supabase(user_email, project_id, report_type, file_name, file_path):
    supabase = get_supabase_client()

    if not supabase or not user_email or not project_id:
        return None

    try:
        payload = {
            "user_email": user_email,
            "project_id": project_id,
            "report_type": report_type,
            "file_name": file_name,
            "file_path": file_path
        }

        result = supabase.table("reports").insert(payload).execute()

        if result.data:
            return result.data[0].get("id")

    except Exception as exc:
        st.warning(f"Supabase report save error: {exc}")

    return None


def load_user_projects(user_email):
    supabase = get_supabase_client()

    if not supabase or not user_email:
        return []

    try:
        result = (
            supabase.table("projects")
            .select("*")
            .eq("user_email", user_email)
            .order("created_at", desc=True)
            .execute()
        )
        return result.data or []
    except Exception as exc:
        st.warning(f"Supabase project load error: {exc}")
        return []


def render_public_landing():
    st.markdown(
        """
<div class="hero">
    <div class="hero-title">OptiFlow Consulting</div>
    <div class="hero-subtitle">
        Endüstri mühendisliği, operasyonel mükemmellik ve veri temelli yönetim raporlaması için profesyonel danışmanlık platformu.
    </div>
    <span class="hero-badge">Operasyonel Verimlilik</span>
    <span class="hero-badge">OEE & KPI Yönetimi</span>
    <span class="hero-badge">Finansal Etki Analizi</span>
    <span class="hero-badge">Yönetim Raporlama</span>
</div>
        """,
        unsafe_allow_html=True
    )

    st.markdown("## Platform Erişimi")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### Demo Sürüm")
        st.write("Örnek veri ile sınırlı demo analizi deneyin. Rapor indirme ve gelişmiş modüller demo sürümde kapalıdır.")
        if st.button("Demo Sürümü Aç", type="secondary"):
            st.session_state.demo_mode = True
            st.rerun()

    with col2:
        st.markdown("### Müşteri Girişi")
        st.write("Kayıtlı müşteriler Google hesabı ile giriş yaparak analiz, rapor ve müşteri portalına erişebilir.")
        if st.button("Müşteri Girişi", type="primary"):
            st.login()

    st.info("Tam platform erişimi için üyelik gerekir. Demo sürüm yalnızca örnek kullanım içindir.")


def render_user_bar(user_email, plan, rules, used, limit):
    st.sidebar.markdown("---")
    st.sidebar.markdown("### Account")

    if st.session_state.get("demo_mode", False):
        st.sidebar.info("Demo Mode")
        if st.sidebar.button("Exit Demo"):
            st.session_state.demo_mode = False
            st.rerun()
    else:
        st.sidebar.success(current_user_name())
        st.sidebar.caption(user_email)
        if st.sidebar.button("Log out"):
            st.logout()

    st.sidebar.markdown(f"**Plan:** {rules.get('label', plan)}")
    st.sidebar.markdown(f"**Usage:** {used}/{limit if limit < 999999 else 'Unlimited'} analyses")

    if plan == "demo":
        st.sidebar.warning("Exports and Danışman Asistanı are locked in Demo.")
    elif plan == "starter":
        st.sidebar.info("Starter: PDF + Excel enabled.")
    elif plan == "professional":
        st.sidebar.success("Professional: Full reporting + enabled.")
    elif plan == "enterprise":
        st.sidebar.success("Enterprise: Full access enabled.")


def render_locked_feature(feature_name, plan):
    st.warning(
        f"{feature_name} is locked for your current plan ({plan}). "
        "Upgrade to unlock this feature."
    )

    starter_link = st.secrets.get("STRIPE_STARTER_LINK", "")
    professional_link = st.secrets.get("STRIPE_PROFESSIONAL_LINK", "")
    enterprise_link = st.secrets.get("ENTERPRISE_CONTACT_LINK", "")

    col1, col2, col3 = st.columns(3)

    with col1:
        if starter_link:
            st.link_button("Upgrade to Starter", starter_link)

    with col2:
        if professional_link:
            st.link_button("Upgrade to Professional", professional_link)

    with col3:
        if enterprise_link:
            st.link_button("Contact Enterprise Sales", enterprise_link)




def create_excel_report(
    company_name,
    sector,
    score,
    maturity,
    company_metrics,
    benchmark_result,
    financial_result,
    recommendations,
    risk_score,
    risk_level
):
    safe_company = (
        str(company_name)
        .replace(" ", "_")
        .replace("/", "_")
        .replace("\\", "_")
        .replace(".", "")
    )

    file_name = os.path.join(
        EXPORT_DIR,
        f"OptiFlow_{safe_company}_Enterprise_Data.xlsx"
    )

    wb = Workbook()

    header_fill = PatternFill("solid", fgColor="0F172A")
    blue_fill = PatternFill("solid", fgColor="EFF6FF")
    green_fill = PatternFill("solid", fgColor="ECFDF5")
    red_fill = PatternFill("solid", fgColor="FFF1F2")
    white_font = Font(color="FFFFFF", bold=True)
    bold_font = Font(bold=True)
    title_font = Font(bold=True, size=14, color="0F172A")
    thin = Side(border_style="thin", color="CBD5E1")

    def style_sheet(ws):
        for row in ws.iter_rows():
            for cell in row:
                cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        for col in ws.columns:
            max_len = 0
            col_letter = col[0].column_letter
            for cell in col:
                try:
                    max_len = max(max_len, len(str(cell.value)))
                except Exception:
                    pass
            ws.column_dimensions[col_letter].width = min(max(max_len + 2, 14), 42)

    def write_table(ws, start_row, title, headers, rows, fill=None):
        ws.cell(start_row, 1, title)
        ws.cell(start_row, 1).font = title_font

        header_row = start_row + 2
        for c, header in enumerate(headers, start=1):
            cell = ws.cell(header_row, c, header)
            cell.fill = header_fill
            cell.font = white_font

        for r, row in enumerate(rows, start=header_row + 1):
            for c, value in enumerate(row, start=1):
                cell = ws.cell(r, c, value)
                if fill:
                    cell.fill = fill

        return header_row + len(rows) + 3

    # DASHBOARD
    ws = wb.active
    ws.title = "Executive Dashboard"

    dashboard_rows = [
        ["Company", company_name],
        ["Sector", sector],
        ["OptiFlow Score", f"{score}/100"],
        ["Maturity", maturity.get("Seviye", "-")],
        ["Risk Level", risk_level],
        ["Risk Score", f"{risk_score}/100"],
        ["Annual Saving", financial_result.get("Tahmini Yıllık Tasarruf", 0)],
        ["ROI (%)", financial_result.get("ROI (%)", 0)],
        ["Payback Month", financial_result.get("Geri Dönüş Süresi (Ay)", 0)],
    ]

    write_table(
        ws,
        1,
        "OptiFlow Consulting Dashboard",
        ["Metric", "Value"],
        dashboard_rows,
        blue_fill
    )

    # KPI
    ws2 = wb.create_sheet("KPI Metrics")
    kpi_rows = [
        ["Wait Rate (%)", company_metrics.get("wait_rate", 0)],
        ["OEE (%)", company_metrics.get("oee", 0)],
        ["Defect Rate (%)", company_metrics.get("defect_rate", 0)],
        ["Line Balance Loss (%)", company_metrics.get("line_balance_loss", 0)],
    ]

    write_table(
        ws2,
        1,
        "Operational KPI Metrics",
        ["KPI", "Value"],
        kpi_rows,
        blue_fill
    )

    # BENCHMARK
    ws3 = wb.create_sheet("Benchmark")
    benchmark_rows = [[str(k), str(v)] for k, v in benchmark_result.items()]

    write_table(
        ws3,
        1,
        "Benchmark Comparison",
        ["Benchmark Field", "Value"],
        benchmark_rows,
        blue_fill
    )

    # FINANCIAL
    ws4 = wb.create_sheet("Financial Impact")
    financial_rows = [[str(k), str(v)] for k, v in financial_result.items()]

    write_table(
        ws4,
        1,
        "Financial Impact Analysis",
        ["Financial Field", "Value"],
        financial_rows,
        green_fill
    )

    # RECOMMENDATIONS
    ws5 = wb.create_sheet("Recommendations")
    recommendation_rows = [[i, rec] for i, rec in enumerate(recommendations, start=1)]

    write_table(
        ws5,
        1,
        "Management Recommendations",
        ["#", "Recommendation"],
        recommendation_rows,
        blue_fill
    )

    # ROADMAP
    ws6 = wb.create_sheet("30-60-90 Roadmap")
    roadmap_rows = [
        ["0-30 Days", "Data validation, bottleneck confirmation, KPI baseline", "Quick-win list and measurement setup"],
        ["30-60 Days", "Flow improvement, line balancing, standard work", "Improved flow and measurable capacity gain"],
        ["60-90 Days", "KPI rhythm, responsibility matrix, performance board", "Sustainable management system"],
    ]

    write_table(
        ws6,
        1,
        "30-60-90 Day Transformation Roadmap",
        ["Period", "Focus", "Expected Output"],
        roadmap_rows,
        blue_fill
    )

    for sheet in wb.worksheets:
        style_sheet(sheet)
        sheet.freeze_panes = "A4"

    wb.save(file_name)
    return file_name


def generate_copilot_answer(
    question,
    company_name,
    sector,
    score,
    maturity,
    company_metrics,
    financial_result,
    recommendations,
    risk_score,
    risk_level
):
    return ask_real_ai_copilot(
        question=question,
        company_name=company_name,
        sector=sector,
        score=score,
        maturity=maturity,
        company_metrics=company_metrics,
        financial_result=financial_result,
        recommendations=recommendations,
        risk_score=risk_score,
        risk_level=risk_level
    )

def render_ai_copilot(
    company_name,
    sector,
    score,
    maturity,
    company_metrics,
    financial_result,
    recommendations,
    risk_score,
    risk_level
):
    st.markdown("## Executive Danışman Asistanı")
    st.write("OpenVeri temelli gerçek yönetim danışmanlığı copilot'u. Mevcut analiz sonuçlarına göre karar desteği üretir.")

    quick_questions = [
        "En büyük operasyonel risk nedir?",
        "ROI ve tasarruf açısından en güçlü fırsat nedir?",
        "İlk 90 gün yol haritası nasıl olmalı?",
        "Öncelikli ilk 3 aksiyon nedir?",
        "Operasyonel olgunluk seviyesi ne ifade ediyor?"
    ]

    selected_question = st.selectbox("Hazır yönetici sorusu seç", quick_questions)

    custom_question = st.text_input(
        "Veya kendi sorunuzu yazın",
        placeholder="Örn: Bu müşteride ilk hangi iyileştirme projesi başlatılmalı?"
    )

    question = custom_question.strip() if custom_question.strip() else selected_question

    if st.button("Copilot Yanıtı Üret", type="primary"):
        answer = generate_copilot_answer(
            question=question,
            company_name=company_name,
            sector=sector,
            score=score,
            maturity=maturity,
            company_metrics=company_metrics,
            financial_result=financial_result,
            recommendations=recommendations,
            risk_score=risk_score,
            risk_level=risk_level
        )

        st.markdown("### Copilot Yanıtı")
        st.markdown(
            f"""
            <div style="padding:22px;border-radius:18px;background:#f8fafc;border-left:6px solid #2563eb;">
                <b>Soru:</b> {question}<br><br>
                <b>Yanıt:</b><br>
                {answer}
            </div>
            """,
            unsafe_allow_html=True
        )

    st.markdown("### Copilot Context")
    st.dataframe(
        pd.DataFrame(
            [
                ["Firma", company_name],
                ["Sektör", sector],
                ["OptiFlow Score", f"{score}/100"],
                ["Risk", f"{risk_level} - {risk_score}/100"],
                ["Yıllık Tasarruf", money_fmt(financial_result.get("Tahmini Yıllık Tasarruf", 0))],
                ["ROI", pct_fmt(financial_result.get("ROI (%)", 0))],
                ["Geri Dönüş", f"{financial_result.get('Geri Dönüş Süresi (Ay)', 0)} Ay"]
            ],
            columns=["Alan", "Değer"]
        ),
        use_container_width=True,
        hide_index=True
    )


# ---------- UI STYLE ----------

st.markdown(
    """
<style>
.hero {
    padding: 34px 34px;
    border-radius: 26px;
    background: linear-gradient(135deg, #0f172a 0%, #1e3a8a 58%, #2563eb 100%);
    color: white;
    margin-bottom: 28px;
}

.hero-title {
    font-size: 44px;
    font-weight: 900;
    margin-bottom: 8px;
}

.hero-subtitle {
    font-size: 17px;
    opacity: 0.92;
    margin-bottom: 18px;
}

.hero-badge {
    display: inline-block;
    padding: 8px 13px;
    border-radius: 999px;
    background-color: rgba(255,255,255,0.14);
    margin-right: 8px;
    font-size: 13px;
    font-weight: 700;
}
</style>
    """,
    unsafe_allow_html=True
)


with st.sidebar:
    if os.path.exists("assets/logo.png"):
        st.image("assets/logo.png", width=220)

    st.title("OptiFlow Portal")

    customer_pages = [
        "Dashboard",
        "Danışman Asistanı",
        "Executive Consultant",
        "Proposal Generator",
        "Analysis",
        "Excel Upload",
        "Benchmark Center",
        "Benchmark Intelligence",
        "Client Portal",
        "My Reports",
        "Report Delivery",
        "Report Center",
        "Billing"
    ]

    admin_pages = [
        "Admin Panel",
        "Admin Analytics",
        "Email Logs"
    ]

    try:
        sidebar_email = st.user.email if st.user.is_logged_in else ""
    except Exception:
        sidebar_email = ""

    admin_email_list = [
        e.strip().lower()
        for e in st.secrets.get("ADMIN_EMAILS", "").split(",")
        if e.strip()
    ]

    nav_pages = customer_pages + (admin_pages if sidebar_email.lower() in admin_email_list else [])

    page = st.radio("Menü", nav_pages)






# ============================================================
# V12 EXCEL UPLOAD + AUTO KPI ENGINE
# ============================================================

def _normalize_col_name(name):
    text = str(name).strip().lower()
    replacements = {
        "ı": "i", "İ": "i", "ğ": "g", "ü": "u", "ş": "s", "ö": "o", "ç": "c"
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace("(", " ").replace(")", " ").replace("/", " ").replace("-", " ").replace("_", " ")
    text = " ".join(text.split())
    return text


def _find_excel_column(df, candidates):
    normalized = {_normalize_col_name(col): col for col in df.columns}

    for candidate in candidates:
        candidate_norm = _normalize_col_name(candidate)
        if candidate_norm in normalized:
            return normalized[candidate_norm]

    for norm_col, original_col in normalized.items():
        for candidate in candidates:
            if _normalize_col_name(candidate) in norm_col:
                return original_col

    return None


def _safe_numeric_series(df, col_name):
    if not col_name or col_name not in df.columns:
        return pd.Series([0] * len(df))

    return pd.to_numeric(df[col_name], errors="coerce").fillna(0)


def auto_kpi_from_uploaded_excel(uploaded_file):
    df = pd.read_excel(uploaded_file)

    if df.empty:
        raise ValueError("Excel dosyası boş görünüyor.")

    process_col = _find_excel_column(df, [
        "process", "süreç", "surec", "operasyon", "istasyon", "station", "workstation", "hat", "line"
    ])

    cycle_col = _find_excel_column(df, [
        "cycle time", "cycle_time", "çevrim süresi", "cevrim suresi", "islem suresi", "işlem süresi",
        "processing time", "operation time", "sure", "süre"
    ])

    wait_col = _find_excel_column(df, [
        "wait minutes", "waiting time", "bekleme", "bekleme süresi", "bekleme suresi",
        "wait", "delay", "gecikme", "queue time"
    ])

    total_qty_col = _find_excel_column(df, [
        "total qty", "total quantity", "production qty", "üretim adedi", "uretim adedi",
        "adet", "quantity", "qty", "output", "uretim", "üretim"
    ])

    good_qty_col = _find_excel_column(df, [
        "good qty", "good quantity", "sağlam", "saglam", "kaliteli adet", "ok qty",
        "accepted qty", "good", "ok"
    ])

    defect_qty_col = _find_excel_column(df, [
        "defect qty", "defect", "scrap", "fire", "hata", "hatalı", "hatali",
        "red", "rejected qty", "reject"
    ])

    planned_time_col = _find_excel_column(df, [
        "planned time", "planned minutes", "planlı süre", "planli sure", "vardiya süresi",
        "vardiya suresi", "available time", "availability time"
    ])

    downtime_col = _find_excel_column(df, [
        "downtime", "down time", "duruş", "durus", "arıza", "ariza", "stop time", "loss time"
    ])

    cycle = _safe_numeric_series(df, cycle_col)
    wait = _safe_numeric_series(df, wait_col)
    total_qty = _safe_numeric_series(df, total_qty_col)
    good_qty = _safe_numeric_series(df, good_qty_col)
    defect_qty = _safe_numeric_series(df, defect_qty_col)
    planned_time = _safe_numeric_series(df, planned_time_col)
    downtime = _safe_numeric_series(df, downtime_col)

    total_cycle = float(cycle.sum())
    total_wait = float(wait.sum())

    if total_cycle + total_wait > 0:
        wait_rate = round((total_wait / (total_cycle + total_wait)) * 100, 1)
    else:
        wait_rate = 0.0

    if defect_qty.sum() > 0 and (good_qty.sum() + defect_qty.sum()) > 0:
        defect_rate = round((float(defect_qty.sum()) / float(good_qty.sum() + defect_qty.sum())) * 100, 1)
    elif total_qty.sum() > 0 and good_qty.sum() > 0:
        defect_rate = round(max(0, (1 - (float(good_qty.sum()) / float(total_qty.sum()))) * 100), 1)
    else:
        defect_rate = 0.0

    if cycle.max() > 0:
        line_balance_loss = round(((float(cycle.max()) - float(cycle.mean())) / float(cycle.max())) * 100, 1)
    else:
        line_balance_loss = 0.0

    if planned_time.sum() > 0:
        availability = max(0, min(1, (float(planned_time.sum()) - float(downtime.sum())) / float(planned_time.sum())))
    else:
        availability = 0.85

    quality = max(0, min(1, 1 - defect_rate / 100))

    if total_qty.sum() > 0 and planned_time.sum() > 0:
        performance = max(0, min(1, float(total_qty.sum()) / max(float(planned_time.sum()), 1)))
        if performance > 1:
            performance = 0.90
    else:
        performance = 0.90

    oee = round(availability * performance * quality * 100, 1)
    if oee <= 0:
        oee = 70.0

    capacity_score = round(max(0, min(100, 100 - line_balance_loss)), 1)

    bottleneck = None
    if process_col and wait_col:
        bottleneck_df = df.copy()
        bottleneck_df["_wait_numeric"] = wait
        bottleneck_row = bottleneck_df.sort_values("_wait_numeric", ascending=False).head(1)
        if not bottleneck_row.empty:
            bottleneck = {
                "process": str(bottleneck_row.iloc[0].get(process_col, "-")),
                "wait_minutes": float(bottleneck_row.iloc[0].get("_wait_numeric", 0))
            }

    pareto_data = None
    if process_col and wait_col:
        pareto_data = (
            df.assign(_wait_numeric=wait)
            .groupby(process_col, dropna=False)["_wait_numeric"]
            .sum()
            .sort_values(ascending=False)
            .reset_index()
            .rename(columns={process_col: "Process", "_wait_numeric": "Wait Minutes"})
        )

    detected_columns = {
        "process": process_col,
        "cycle_time": cycle_col,
        "wait_minutes": wait_col,
        "total_qty": total_qty_col,
        "good_qty": good_qty_col,
        "defect_qty": defect_qty_col,
        "planned_time": planned_time_col,
        "downtime": downtime_col
    }

    metrics = {
        "wait_rate": wait_rate,
        "oee": oee,
        "defect_rate": defect_rate,
        "line_balance_loss": line_balance_loss,
        "capacity_score": capacity_score
    }

    financial_inputs = {
        "total_wait_minutes": round(total_wait, 1) if total_wait > 0 else 120,
        "hourly_labor_cost": 250,
        "working_days": 22,
        "improvement_rate": 20
    }

    return {
        "dataframe": df,
        "metrics": metrics,
        "financial_inputs": financial_inputs,
        "detected_columns": detected_columns,
        "bottleneck": bottleneck,
        "pareto_data": pareto_data
    }


def create_v12_excel_template():
    file_name = os.path.join(EXPORT_DIR, "OptiFlow_V12_Upload_Template.xlsx")

    wb = Workbook()
    ws = wb.active
    ws.title = "Production Data"

    headers = [
        "Process",
        "Cycle Time",
        "Wait Minutes",
        "Production Qty",
        "Good Qty",
        "Defect Qty",
        "Planned Time",
        "Downtime"
    ]

    ws.append(headers)

    sample_rows = [
        ["Cutting", 45, 18, 120, 116, 4, 480, 20],
        ["Sewing", 62, 35, 110, 104, 6, 480, 35],
        ["Quality Control", 38, 12, 105, 101, 4, 480, 15],
        ["Packing", 30, 8, 100, 98, 2, 480, 10],
    ]

    for row in sample_rows:
        ws.append(row)

    for col in ws.columns:
        col_letter = col[0].column_letter
        ws.column_dimensions[col_letter].width = 18
        for cell in col:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    wb.save(file_name)
    return file_name


def render_v12_excel_upload_center(
    company_name,
    sector,
    active_plan,
    active_rules
):
    st.markdown("## Excel Upload + Otomatik KPI Motoru")
    st.write("Üretim veya operasyon verisini yükle; OptiFlow otomatik KPI, darboğaz, benchmark ve finansal analiz üretir.")

    if active_plan == "demo":
        st.warning("Demo planda Excel Upload sınırlıdır. Örnek şablonu indirip test edebilirsin; rapor export kilitlidir.")

    template_path = create_v12_excel_template()
    with open(template_path, "rb") as file:
        st.download_button(
            "OptiFlow Excel Şablonu İndir",
            data=file,
            file_name="OptiFlow_V12_Upload_Template.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    uploaded_file = st.file_uploader(
        "Excel dosyası yükle",
        type=["xlsx", "xls"],
        help="Önerilen kolonlar: Process, Cycle Time, Wait Minutes, Production Qty, Good Qty, Defect Qty, Planned Time, Downtime"
    )

    if not uploaded_file:
        st.info("Excel dosyası yüklediğinde otomatik KPI hesaplama başlayacak.")
        return

    try:
        result = auto_kpi_from_uploaded_excel(uploaded_file)
    except Exception as exc:
        st.error(f"Excel okunamadı: {exc}")
        return

    metrics = result["metrics"]
    financial_inputs = result["financial_inputs"]

    st.success("Excel başarıyla analiz edildi.")

    st.markdown("### Algılanan Kolonlar")
    st.json(result["detected_columns"])

    st.markdown("### Otomatik KPI Sonuçları")
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Bekleme Oranı", f"%{metrics['wait_rate']}")
    c2.metric("OEE", f"%{metrics['oee']}")
    c3.metric("Hata Oranı", f"%{metrics['defect_rate']}")
    c4.metric("Hat Denge Kaybı", f"%{metrics['line_balance_loss']}")
    c5.metric("Kapasite Skoru", f"%{metrics['capacity_score']}")

    if result.get("bottleneck"):
        st.warning(
            f"Kritik darboğaz: {result['bottleneck']['process']} - "
            f"{result['bottleneck']['wait_minutes']} dk bekleme"
        )

    st.markdown("### Yüklenen Veri Önizleme")
    st.dataframe(result["dataframe"], use_container_width=True)

    if result.get("pareto_data") is not None:
        st.markdown("### Bekleme Pareto Analizi")
        st.dataframe(result["pareto_data"], use_container_width=True)
        try:
            st.bar_chart(result["pareto_data"].set_index("Process"), use_container_width=True)
        except Exception:
            pass

    auto_company_metrics = {
        "wait_rate": metrics["wait_rate"],
        "oee": metrics["oee"],
        "defect_rate": metrics["defect_rate"],
        "line_balance_loss": metrics["line_balance_loss"]
    }

    auto_benchmark = benchmark_from_database(auto_company_metrics, sector)

    auto_score = calculate_optiflow_score(
        efficiency_score=100 - metrics["wait_rate"],
        oee_score=metrics["oee"],
        capacity_score=metrics["capacity_score"],
        flow_score=100 - metrics["line_balance_loss"],
        quality_score=100 - metrics["defect_rate"]
    )

    auto_maturity = get_maturity_comment(auto_score)

    auto_financial = calculate_financial_impact(
        total_wait_minutes=financial_inputs["total_wait_minutes"],
        improvement_rate=financial_inputs["improvement_rate"],
        hourly_labor_cost=financial_inputs["hourly_labor_cost"],
        working_days_per_month=financial_inputs["working_days"]
    )

    auto_recommendations = generate_recommendations(
        wait_rate=metrics["wait_rate"],
        oee=metrics["oee"],
        line_balance_loss=metrics["line_balance_loss"]
    )

    auto_risk_score, auto_risk_level, _, _ = risk_status(
        metrics["wait_rate"],
        metrics["oee"],
        metrics["defect_rate"],
        metrics["line_balance_loss"]
    )

    st.markdown("### Otomatik OptiFlow Değerlendirmesi")
    render_plotly_dashboard(
        company_name=company_name,
        score=auto_score,
        maturity=auto_maturity,
        company_metrics=auto_company_metrics,
        benchmark_result=auto_benchmark,
        financial_result=auto_financial,
        risk_score=auto_risk_score,
        risk_level=auto_risk_level
    )

    st.markdown("### V12 Rapor Üretimi")

    if not active_rules.get("pdf", False):
        render_locked_feature("V12 PDF / Excel Export", active_plan)
        return

    if st.button("Excel Verisinden PDF + Excel Rapor Oluştur", type="primary"):
        with st.spinner("V12 otomatik rapor hazırlanıyor..."):
            consulting_report = generate_consulting_report(
                sector=sector,
                company_metrics=auto_company_metrics,
                benchmark_result=auto_benchmark,
                financial_result=auto_financial,
                maturity=auto_maturity,
                recommendations=auto_recommendations
            )

            pdf_file = create_enterprise_pdf(
                company_name=company_name,
                sector=sector,
                score=auto_score,
                maturity=auto_maturity,
                company_metrics=auto_company_metrics,
                benchmark_result=auto_benchmark,
                financial_result=auto_financial,
                recommendations=auto_recommendations,
                consulting_report=consulting_report
            )

            excel_file = create_excel_report(
                company_name=company_name,
                sector=sector,
                score=auto_score,
                maturity=auto_maturity,
                company_metrics=auto_company_metrics,
                benchmark_result=auto_benchmark,
                financial_result=auto_financial,
                recommendations=auto_recommendations,
                risk_score=auto_risk_score,
                risk_level=auto_risk_level
            )

            project_id = save_project_to_supabase(
                user_email=user_email,
                company_name=company_name,
                sector=sector,
                score=auto_score,
                risk_level=auto_risk_level,
                financial_result=auto_financial
            )

            if project_id:
                save_report_to_supabase(user_email, project_id, "pdf", os.path.basename(pdf_file), pdf_file)
                save_report_to_supabase(user_email, project_id, "excel", os.path.basename(excel_file), excel_file)

        st.success("V12 raporları oluşturuldu.")

        with open(pdf_file, "rb") as file:
            st.download_button(
                "V12 Enterprise PDF İndir",
                data=file,
                file_name=f"OptiFlow_{company_name.replace(' ', '_')}_V12_Report.pdf",
                mime="application/pdf"
            )

        with open(excel_file, "rb") as file:
            st.download_button(
                "V12 Excel Data İndir",
                data=file,
                file_name=os.path.basename(excel_file),
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )



# ============================================================
# ADMIN PANEL LAYER
# ============================================================

def get_admin_emails():
    raw = st.secrets.get("ADMIN_EMAILS", "")
    if not raw:
        return []
    return [email.strip().lower() for email in raw.split(",") if email.strip()]


def is_admin_user(email):
    return str(email or "").lower() in get_admin_emails()


def admin_load_table(table_name):
    supabase = get_supabase_client()
    if not supabase:
        return []

    try:
        result = supabase.table(table_name).select("*").execute()
        return result.data or []
    except Exception as exc:
        st.error(f"Admin table load error ({table_name}): {exc}")
        return []


def admin_update_user_plan(email, plan, status):
    supabase = get_supabase_client()
    if not supabase:
        return False

    rules = PLAN_RULES.get(plan, PLAN_RULES["demo"])

    try:
        supabase.table("users").update({
            "plan": plan,
            "plan_status": status,
            "analysis_limit": int(rules.get("analysis_limit", 1))
        }).eq("email", email).execute()

        supabase.table("subscriptions").upsert({
            "user_email": email,
            "plan": plan,
            "status": status
        }).execute()

        return True
    except Exception as exc:
        st.error(f"Plan update error: {exc}")
        return False


def render_admin_dashboard(user_email):
    if not is_admin_user(user_email):
        st.error("Bu sayfa sadece admin kullanıcılar içindir.")
        st.stop()

    st.markdown("## Admin Dashboard")
    st.caption("OptiFlow Platform yönetim paneli")

    users = admin_load_table("users")
    projects = admin_load_table("projects")
    reports = admin_load_table("reports")
    subscriptions = admin_load_table("subscriptions")

    total_users = len(users)
    total_projects = len(projects)
    total_reports = len(reports)
    active_subs = len([s for s in subscriptions if str(s.get("status", "")).lower() in ["active", "trialing"]])
    demo_users = len([u for u in users if str(u.get("plan", "demo")).lower() == "demo"])
    professional_users = len([u for u in users if str(u.get("plan", "")).lower() == "professional"])
    enterprise_users = len([u for u in users if str(u.get("plan", "")).lower() == "enterprise"])

    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Toplam Kullanıcı", total_users)
    c2.metric("Aktif Abone", active_subs)
    c3.metric("Demo Kullanıcı", demo_users)
    c4.metric("Toplam Analiz", total_projects)
    c5.metric("Toplam Rapor", total_reports)

    st.markdown("### Plan Dağılımı")
    plan_rows = []
    for plan in ["demo", "starter", "professional", "enterprise"]:
        count = len([u for u in users if str(u.get("plan", "demo")).lower() == plan])
        plan_rows.append({"Plan": plan, "Kullanıcı Sayısı": count})

    st.dataframe(plan_rows, use_container_width=True, hide_index=True)

    st.markdown("### Kullanıcı Yönetimi")

    if not users:
        st.info("Henüz kullanıcı yok.")
    else:
        user_options = [u.get("email") for u in users if u.get("email")]
        selected_email = st.selectbox("Kullanıcı seç", user_options)

        selected_user = next((u for u in users if u.get("email") == selected_email), {})
        st.json(selected_user)

        col_plan, col_status, col_button = st.columns([2, 2, 1])

        with col_plan:
            new_plan = st.selectbox(
                "Plan",
                ["demo", "starter", "professional", "enterprise"],
                index=["demo", "starter", "professional", "enterprise"].index(
                    str(selected_user.get("plan", "demo")).lower()
                    if str(selected_user.get("plan", "demo")).lower() in ["demo", "starter", "professional", "enterprise"]
                    else "demo"
                )
            )

        with col_status:
            new_status = st.selectbox(
                "Status",
                ["inactive", "active", "trialing", "canceled"],
                index=1 if str(selected_user.get("plan_status", "inactive")).lower() == "active" else 0
            )

        with col_button:
            st.write("")
            st.write("")
            if st.button("Güncelle", type="primary"):
                ok = admin_update_user_plan(selected_email, new_plan, new_status)
                if ok:
                    st.success("Kullanıcı planı güncellendi.")
                    st.rerun()

    st.markdown("### Son Analizler")
    if projects:
        project_df = pd.DataFrame(projects)
        columns = [c for c in ["user_email", "company_name", "sector", "score", "risk_level", "annual_saving", "created_at"] if c in project_df.columns]
        st.dataframe(project_df[columns], use_container_width=True, hide_index=True)
    else:
        st.info("Henüz analiz kaydı yok.")

    st.markdown("### Rapor Kayıtları")
    if reports:
        report_df = pd.DataFrame(reports)
        columns = [c for c in ["user_email", "report_type", "file_name", "created_at"] if c in report_df.columns]
        st.dataframe(report_df[columns], use_container_width=True, hide_index=True)
    else:
        st.info("Henüz rapor kaydı yok.")











# ============================================================
# V17 CLIENT PORTAL + V18 BENCHMARK INTELLIGENCE + V19 PROPOSAL GENERATOR
# ============================================================

def render_client_portal(user_email):
    st.markdown("## Client Portal")
    st.caption("Müşteri bazlı rapor geçmişi, KPI trendleri, tasarruf potansiyeli ve aksiyon görünümü")

    projects = load_user_projects(user_email)
    reports = load_user_reports(user_email)

    if not projects:
        st.info("Henüz proje kaydı yok. Report Center veya Excel Upload üzerinden rapor oluşturduğunda burada görünecek.")
        return

    projects_df = pd.DataFrame(projects)

    if "created_at" in projects_df.columns:
        projects_df["created_at"] = pd.to_datetime(projects_df["created_at"], errors="coerce")

    total_projects = len(projects_df)
    avg_score = pd.to_numeric(projects_df.get("score", 0), errors="coerce").fillna(0).mean()
    total_saving = pd.to_numeric(projects_df.get("annual_saving", 0), errors="coerce").fillna(0).sum()

    high_risk_count = 0
    if "risk_level" in projects_df.columns:
        high_risk_count = int(projects_df["risk_level"].astype(str).str.lower().isin(["yüksek", "high"]).sum())

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Toplam Analiz", total_projects)
    c2.metric("Ortalama Score", f"{avg_score:.1f}")
    c3.metric("Toplam Tasarruf Potansiyeli", money_fmt(total_saving))
    c4.metric("Yüksek Riskli Analiz", high_risk_count)

    st.markdown("### KPI Trendleri")

    trend_cols = st.columns(2)

    with trend_cols[0]:
        if "created_at" in projects_df.columns and "score" in projects_df.columns:
            temp = projects_df.dropna(subset=["created_at"]).sort_values("created_at")
            if not temp.empty:
                fig = px.line(
                    temp,
                    x="created_at",
                    y="score",
                    color="company_name" if "company_name" in temp.columns else None,
                    markers=True,
                    title="OptiFlow Score Trend"
                )
                fig.update_layout(height=360, margin=dict(l=20, r=20, t=55, b=20))
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.info("Trend için tarih verisi yok.")
        else:
            st.info("Score trend verisi yok.")

    with trend_cols[1]:
        if "created_at" in projects_df.columns and "annual_saving" in projects_df.columns:
            temp = projects_df.dropna(subset=["created_at"]).sort_values("created_at")
            temp["annual_saving"] = pd.to_numeric(temp["annual_saving"], errors="coerce").fillna(0)
            if not temp.empty:
                fig = px.bar(
                    temp,
                    x="created_at",
                    y="annual_saving",
                    color="company_name" if "company_name" in temp.columns else None,
                    title="Tasarruf Potansiyeli Trend"
                )
                fig.update_layout(height=360, margin=dict(l=20, r=20, t=55, b=20))
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.info("Tasarruf trend verisi yok.")
        else:
            st.info("Tasarruf trend verisi yok.")

    st.markdown("### Proje Geçmişi")
    visible_cols = [c for c in ["company_name", "sector", "score", "risk_level", "annual_saving", "created_at"] if c in projects_df.columns]
    st.dataframe(projects_df[visible_cols], use_container_width=True, hide_index=True)

    st.markdown("### Rapor Geçmişi")
    if reports:
        reports_df = pd.DataFrame(reports)
        report_cols = [c for c in ["report_type", "file_name", "created_at"] if c in reports_df.columns]
        st.dataframe(reports_df[report_cols], use_container_width=True, hide_index=True)
    else:
        st.info("Henüz rapor kaydı yok.")

    st.markdown("### Açık Yönetim Aksiyonları")

    action_rows = []
    latest = projects_df.sort_values("created_at", ascending=False).head(5) if "created_at" in projects_df.columns else projects_df.head(5)

    for _, row in latest.iterrows():
        risk = str(row.get("risk_level", "")).lower()
        score_val = float(row.get("score", 0) or 0)

        if risk in ["yüksek", "high"] or score_val < 65:
            priority = "Yüksek"
            action = "Darboğaz, bekleme süresi ve hat dengeleme çalışması başlatılmalı."
        elif score_val < 80:
            priority = "Orta"
            action = "KPI ritmi, standart iş ve OEE takip sistemi güçlendirilmeli."
        else:
            priority = "Düşük"
            action = "Sürdürülebilir performans yönetimi ve iyileştirme fırsatları izlenmeli."

        action_rows.append({
            "Firma": row.get("company_name", "-"),
            "Öncelik": priority,
            "Aksiyon": action,
            "Durum": "Açık"
        })

    if action_rows:
        st.dataframe(pd.DataFrame(action_rows), use_container_width=True, hide_index=True)
    else:
        st.info("Aksiyon oluşturmak için proje kaydı yok.")


def benchmark_intelligence_text(company_metrics, benchmark_result, sector, score, risk_level):
    insights = []

    metric_pairs = [
        ("Bekleme Oranı", "Firma Bekleme Oranı", "Sektör Bekleme Oranı", "lower"),
        ("OEE", "Firma OEE", "Sektör OEE", "higher"),
        ("Hata Oranı", "Firma Hata Oranı", "Sektör Hata Oranı", "lower"),
        ("Hat Denge Kaybı", "Firma Hat Denge Kaybı", "Sektör Hat Denge Kaybı", "lower"),
    ]

    for label, firm_key, sector_key, better in metric_pairs:
        try:
            firm_val = float(benchmark_result.get(firm_key, 0))
            sector_val = float(benchmark_result.get(sector_key, 0))
            diff = firm_val - sector_val

            if sector_val == 0:
                continue

            diff_pct = abs(diff / sector_val * 100)

            if better == "lower":
                if firm_val <= sector_val:
                    insights.append(f"{label}: Firma sektör ortalamasından %{diff_pct:.1f} daha iyi konumda.")
                else:
                    insights.append(f"{label}: Firma sektör ortalamasından %{diff_pct:.1f} daha zayıf; iyileştirme önceliği yüksek.")
            else:
                if firm_val >= sector_val:
                    insights.append(f"{label}: Firma sektör ortalamasından %{diff_pct:.1f} daha güçlü performans gösteriyor.")
                else:
                    insights.append(f"{label}: Firma sektör ortalamasının %{diff_pct:.1f} gerisinde; kapasite ve etkinlik iyileştirmesi gerekli.")
        except Exception:
            pass

    if score >= 80:
        position = "güçlü ve ölçeklenebilir operasyonel yapı"
    elif score >= 65:
        position = "yönetilebilir fakat iyileştirme potansiyeli yüksek operasyonel yapı"
    else:
        position = "yüksek dönüşüm ihtiyacı olan operasyonel yapı"

    executive = (
        f"{sector} sektöründe yapılan benchmark değerlendirmesine göre şirketin genel pozisyonu "
        f"'{position}' olarak yorumlanabilir. OptiFlow Score {score}/100, risk seviyesi ise {risk_level}. "
        "Öncelik; sektör ortalamasından negatif ayrışan KPI'ların finansal etki sırasına göre iyileştirilmesidir."
    )

    return executive, insights


def render_benchmark_intelligence(
    company_name,
    sector,
    score,
    company_metrics,
    benchmark_result,
    financial_result,
    risk_level
):
    st.markdown("## V18 Benchmark Intelligence")
    st.caption("Sektör karşılaştırmasını yönetici diline çeviren stratejik benchmark motoru")

    executive, insights = benchmark_intelligence_text(
        company_metrics=company_metrics,
        benchmark_result=benchmark_result,
        sector=sector,
        score=score,
        risk_level=risk_level
    )

    st.markdown(
        f"""
        <div style="padding:22px;border-radius:18px;background:#eff6ff;border-left:6px solid #2563eb;">
        <b>Executive Benchmark Insight</b><br><br>
        {executive}
        </div>
        """,
        unsafe_allow_html=True
    )

    st.markdown("### KPI Intelligence")

    if insights:
        for item in insights:
            st.markdown(f"- {item}")
    else:
        st.info("Benchmark içgörüsü üretmek için yeterli veri bulunamadı.")

    st.markdown("### Firma vs Sektör Karşılaştırması")
    st.plotly_chart(plot_benchmark(benchmark_result), use_container_width=True)

    st.markdown("### Öncelik Matrisi")

    priority_rows = []
    for item in insights:
        if "zayıf" in item.lower() or "gerisinde" in item.lower():
            priority_rows.append(["Yüksek", item, "İlk 30 gün içinde aksiyon planına alınmalı"])
        else:
            priority_rows.append(["Orta/Düşük", item, "Mevcut uygulama korunmalı ve izlenmeli"])

    if priority_rows:
        st.dataframe(
            pd.DataFrame(priority_rows, columns=["Öncelik", "Benchmark Bulgusu", "Yönetim Aksiyonu"]),
            use_container_width=True,
            hide_index=True
        )

    st.markdown("### Benchmark Yorumu")

    if st.button("Benchmark Intelligence Raporu Üret", type="primary"):
        question = f"""
Aşağıdaki benchmark sonuçlarına göre üst yönetime yönelik benchmark intelligence raporu yaz:

Firma: {company_name}
Sektör: {sector}
OptiFlow Score: {score}/100
Risk Seviyesi: {risk_level}
Finansal Sonuçlar: {financial_result}
Benchmark Sonuçları: {benchmark_result}

Çıktı formatı:
1. Sektör Pozisyonu
2. Kritik Ayrışmalar
3. Finansal Öncelik
4. Operasyonel Öncelik
5. Yönetim Kararı
"""
        with st.spinner("Benchmark Intelligence hazırlanıyor..."):
            answer = ask_real_ai_copilot(
                question=question,
                company_name=company_name,
                sector=sector,
                score=score,
                maturity={"Seviye": "-", "Yorum": "-"},
                company_metrics=company_metrics,
                financial_result=financial_result,
                recommendations=[],
                risk_score=0,
                risk_level=risk_level
            )

        st.markdown(
            f"""
            <div style="padding:22px;border-radius:18px;background:#f8fafc;border-left:6px solid #0f172a;">
            {answer.replace(chr(10), '<br>')}
            </div>
            """,
            unsafe_allow_html=True
        )


def generate_consulting_proposal_text(
    company_name,
    sector,
    score,
    maturity,
    company_metrics,
    financial_result,
    recommendations,
    risk_level
):
    yearly_saving = financial_result.get("Tahmini Yıllık Tasarruf", 0)
    roi = financial_result.get("ROI (%)", 0)
    payback = financial_result.get("Geri Dönüş Süresi (Ay)", 0)

    proposal = f"""
# OptiFlow Consulting Proposal

## 1. Teklif Özeti

{company_name} için gerçekleştirilen ön değerlendirme, işletmenin operasyonel performansında ölçülebilir iyileştirme potansiyeli bulunduğunu göstermektedir. OptiFlow Score {score}/100, operasyonel olgunluk seviyesi {maturity.get("Seviye", "-")} ve risk seviyesi {risk_level} olarak değerlendirilmiştir.

Bu teklifin amacı; bekleme süreleri, hat dengeleme kayıpları, OEE performansı, kalite kayıpları ve KPI yönetim sistemini iyileştirerek finansal ve operasyonel değer yaratmaktır.

## 2. Mevcut Durum Bulguları

- Bekleme Oranı: %{company_metrics.get("wait_rate", 0)}
- OEE: %{company_metrics.get("oee", 0)}
- Hata Oranı: %{company_metrics.get("defect_rate", 0)}
- Hat Denge Kaybı: %{company_metrics.get("line_balance_loss", 0)}
- Tahmini Yıllık Tasarruf Potansiyeli: {money_fmt(yearly_saving)}
- ROI: {pct_fmt(roi)}
- Geri Dönüş Süresi: {payback} ay

## 3. Proje Kapsamı

OptiFlow danışmanlık projesi aşağıdaki iş paketlerini kapsar:

### İş Paketi 1 - Veri Doğrulama ve Operasyonel Teşhis
Süreç gözlemi, KPI baz çizgisinin doğrulanması, darboğaz noktalarının tespiti ve veri güvenilirliğinin değerlendirilmesi.

### İş Paketi 2 - Akış ve Hat Dengeleme İyileştirmesi
Bekleme sürelerinin azaltılması, iş yükü dağılımının dengelenmesi, standart iş uygulamalarının tasarlanması.

### İş Paketi 3 - OEE ve Performans Yönetim Sistemi
OEE bileşen takibi, duruş neden analizi, günlük KPI panosu ve haftalık yönetim ritminin kurulması.

### İş Paketi 4 - Finansal Etki ve ROI Takibi
Operasyonel kayıpların parasal etkisinin izlenmesi, iyileştirme aksiyonlarının finansal karşılığının raporlanması.

## 4. 30-60-90 Gün Proje Planı

### 0-30 Gün
Veri doğrulama, süreç gözlemi, darboğaz analizi, KPI baz çizgisi ve hızlı kazanım alanlarının belirlenmesi.

### 30-60 Gün
Hat dengeleme, bekleme azaltma, standart iş tasarımı, OEE takip altyapısı ve pilot uygulamalar.

### 60-90 Gün
KPI toplantı ritmi, sorumluluk matrisi, performans panosu, sürdürülebilir yönetim sistemi ve sonuç raporu.

## 5. Beklenen Çıktılar

- Ölçülebilir KPI iyileşmesi
- Bekleme ve akış kayıplarında azalma
- Daha dengeli kapasite kullanımı
- Kalite ve yeniden işleme kayıplarında görünürlük
- Yönetim için veri temelli karar destek sistemi
- PDF, PPT ve Excel formatında yönetim raporları

## 6. Öncelikli Aksiyonlar

{chr(10).join([f"- {rec}" for rec in recommendations])}

## 7. Ticari Yaklaşım

Önerilen proje yaklaşımı üç fazlıdır:

- Faz 1: Operasyonel Teşhis ve Veri Doğrulama
- Faz 2: İyileştirme Tasarımı ve Pilot Uygulama
- Faz 3: Sürdürülebilir Yönetim Sistemi ve Raporlama

## 8. Yönetim Kararı Önerisi

{company_name} için önerilen ilk adım, 30 günlük operasyonel teşhis ve hızlı kazanım çalışmasının başlatılmasıdır. Bu çalışma sonrasında daha kapsamlı 90 günlük dönüşüm programı netleştirilebilir.
"""
    return proposal


def create_proposal_docx(company_name, proposal_text):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    safe_company = str(company_name).replace(" ", "_").replace("/", "_")
    file_name = os.path.join(EXPORT_DIR, f"OptiFlow_{safe_company}_Consulting_Proposal.docx")

    doc = Document()
    title = doc.add_heading("OptiFlow Consulting Proposal", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in str(proposal_text).split("\n"):
        clean = line.strip()
        if not clean:
            doc.add_paragraph("")
            continue

        if clean.startswith("# "):
            doc.add_heading(clean.replace("# ", ""), level=1)
        elif clean.startswith("## "):
            doc.add_heading(clean.replace("## ", ""), level=2)
        elif clean.startswith("### "):
            doc.add_heading(clean.replace("### ", ""), level=3)
        else:
            p = doc.add_paragraph(clean)
            for run in p.runs:
                run.font.size = Pt(10.5)

    doc.save(file_name)
    return file_name


def render_proposal_generator(
    company_name,
    sector,
    score,
    maturity,
    company_metrics,
    financial_result,
    recommendations,
    risk_level,
    active_plan,
    active_rules
):
    st.markdown("## Proposal Generator")
    st.caption("OptiFlow analizinden otomatik danışmanlık teklif dokümanı üretir.")

    if not active_rules.get("ai", False):
        render_locked_feature("Proposal Generator", active_plan)
        return

    st.markdown("### Teklif Parametreleri")

    col1, col2, col3 = st.columns(3)
    with col1:
        proposal_type = st.selectbox("Teklif Tipi", ["Operational Excellence", "Lean Transformation", "OEE Improvement", "KPI Management System"])
    with col2:
        project_duration = st.selectbox("Proje Süresi", ["30 Gün", "60 Gün", "90 Gün", "6 Ay"])
    with col3:
        commercial_model = st.selectbox("Ticari Model", ["Sabit Proje Bedeli", "Aylık Danışmanlık", "Başarı Bazlı", "Hibrit Model"])

    if "v19_proposal_text" not in st.session_state:
        st.session_state.v19_proposal_text = ""

    if st.button("Danışmanlık Teklifi Oluştur", type="primary"):
        with st.spinner("OptiFlow danışmanlık teklifi hazırlanıyor..."):
            base_proposal = generate_consulting_proposal_text(
                company_name=company_name,
                sector=sector,
                score=score,
                maturity=maturity,
                company_metrics=company_metrics,
                financial_result=financial_result,
                recommendations=recommendations,
                risk_level=risk_level
            )

            ai_question = f"""
Aşağıdaki danışmanlık teklif taslağını daha profesyonel, satışa uygun ve yönetim danışmanlığı diliyle geliştir.

Teklif tipi: {proposal_type}
Proje süresi: {project_duration}
Ticari model: {commercial_model}

Taslak:
{base_proposal}

Çıktı Türkçe olsun.
Başlıklar korunmalı.
Danışmanlık şirketi müşteriye gönderecekmiş gibi yaz.
"""
            enhanced = ask_real_ai_copilot(
                question=ai_question,
                company_name=company_name,
                sector=sector,
                score=score,
                maturity=maturity,
                company_metrics=company_metrics,
                financial_result=financial_result,
                recommendations=recommendations,
                risk_score=0,
                risk_level=risk_level
            )

            st.session_state.v19_proposal_text = enhanced

    if st.session_state.v19_proposal_text:
        st.markdown("### Proposal Output")
        st.markdown(
            f"""
            <div style="padding:24px;border-radius:18px;background:#f8fafc;border-left:6px solid #2563eb;">
            {st.session_state.v19_proposal_text.replace(chr(10), '<br>')}
            </div>
            """,
            unsafe_allow_html=True
        )

        proposal_docx = create_proposal_docx(company_name, st.session_state.v19_proposal_text)

        with open(proposal_docx, "rb") as file:
            st.download_button(
                "Consulting Proposal DOCX İndir",
                data=file,
                file_name=os.path.basename(proposal_docx),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        st.markdown("### Teklifi Mail Gönder")
        proposal_to = st.text_input("Alıcı e-posta", value=user_email or "", key="proposal_mail_to")

        if st.button("Proposal Mail Gönder"):
            ok, msg = send_email_with_attachments(
                to_email=proposal_to,
                subject=f"OptiFlow Consulting Proposal - {company_name}",
                body=f"""Merhaba,

{company_name} için hazırlanan OptiFlow danışmanlık teklif dokümanı ekte yer almaktadır.

Saygılarımızla,
OptiFlow Consulting""",
                attachment_paths=[proposal_docx]
            )

            save_email_log(
                user_email=user_email,
                to_email=proposal_to,
                subject=f"OptiFlow Consulting Proposal - {company_name}",
                status="success" if ok else "failed",
                message=msg,
                company_name=company_name
            )

            if ok:
                st.success("Teklif e-posta ile gönderildi.")
            else:
                st.error(msg)


# ============================================================
# EXECUTIVE CONSULTANT
# ============================================================

def build_executive_consultant_context(
    company_name,
    sector,
    score,
    maturity,
    company_metrics,
    benchmark_result,
    financial_result,
    recommendations,
    risk_score,
    risk_level
):
    return f"""
Company: {company_name}
Sector: {sector}

OptiFlow Score: {score}/100
Maturity Level: {maturity.get("Seviye", "-")}
Maturity Comment: {maturity.get("Yorum", "-")}

Operational KPIs:
- Wait Rate: {company_metrics.get("wait_rate", 0)}%
- OEE: {company_metrics.get("oee", 0)}%
- Defect Rate: {company_metrics.get("defect_rate", 0)}%
- Line Balance Loss: {company_metrics.get("line_balance_loss", 0)}%

Risk:
- Risk Score: {risk_score}/100
- Risk Level: {risk_level}

Financial:
- Total Operational Loss: {financial_result.get("Toplam Operasyonel Kayıp", 0)} TL
- Improvement Potential: {financial_result.get("İyileştirme Potansiyeli", 0)} TL
- Estimated Annual Saving: {financial_result.get("Tahmini Yıllık Tasarruf", 0)} TL
- ROI: {financial_result.get("ROI (%)", 0)}%
- Payback Period: {financial_result.get("Geri Dönüş Süresi (Ay)", 0)} months

Benchmark:
{benchmark_result}

Recommendations:
{chr(10).join([f"- {rec}" for rec in recommendations])}
"""


def generate_executive_consultant_output(
    mode,
    company_name,
    sector,
    score,
    maturity,
    company_metrics,
    benchmark_result,
    financial_result,
    recommendations,
    risk_score,
    risk_level
):
    context = build_executive_consultant_context(
        company_name=company_name,
        sector=sector,
        score=score,
        maturity=maturity,
        company_metrics=company_metrics,
        benchmark_result=benchmark_result,
        financial_result=financial_result,
        recommendations=recommendations,
        risk_score=risk_score,
        risk_level=risk_level
    )

    mode_prompts = {
        "CEO Summary": """
Sen üst düzey bir yönetim danışmanısın. CEO'ya sunulacak 1 sayfalık yönetici özeti üret.
Odak: stratejik öncelikler, risk, finansal fırsat, hızlı kazanımlar, karar önerisi.
""",
        "COO Operations Brief": """
Sen operasyon direktörüne danışmanlık yapan kıdemli operasyonel mükemmellik uzmanısın.
Odak: bekleme, OEE, hat dengeleme, darboğaz, kalite, standart iş, KPI yönetim ritmi.
""",
        "CFO Financial Brief": """
Sen CFO'ya rapor hazırlayan finansal etki danışmanısın.
Odak: kayıp, tasarruf potansiyeli, ROI, geri dönüş süresi, yatırım önceliği, finansal risk.
""",
        "90-Day Transformation Plan": """
Sen dönüşüm programı yöneten senior consultant'sın.
30-60-90 gün için ayrıntılı ama uygulanabilir aksiyon planı hazırla.
Her faz için hedef, faaliyet, sorumlu ekip, beklenen çıktı, KPI yaz.
""",
        "Board Presentation Notes": """
Sen yönetim kurulu sunumu hazırlıyorsun.
Kısa, güçlü, ikna edici ve sayısal veri temelli board meeting konuşma notları üret.
""",
        "Sales Proposal": """
Sen danışmanlık şirketi için müşteri teklif metni hazırlıyorsun.
OptiFlow analizine göre hizmet kapsamı, proje yaklaşımı, beklenen değer ve sonraki adımları yaz.
"""
    }

    instruction = mode_prompts.get(mode, mode_prompts["CEO Summary"])

    prompt = f"""
{instruction}

Aşağıdaki OptiFlow analiz verilerini kullan:

{context}

Çıktıyı Türkçe yaz.
Profesyonel yönetim danışmanlığı dili kullan.
Gereksiz uzun yazma ama somut ve sayısal ol.
Başlıklar net olsun.
Cümleler satışa ve yönetime uygun olsun.

Format:
1. Yönetici Özeti
2. Kritik Bulgular
3. Finansal Etki
4. Risk ve Önceliklendirme
5. İlk 3 Aksiyon
6. 30-60-90 Gün Yol Haritası
7. Yönetim Kararı Önerisi
"""

    try:
        # Mevcut ai_copilot modülünü kullanmak yerine doğrudan Openfonksiyonuna uygun soru soruyoruz.
        answer = ask_real_ai_copilot(
            question=prompt,
            company_name=company_name,
            sector=sector,
            score=score,
            maturity=maturity,
            company_metrics=company_metrics,
            financial_result=financial_result,
            recommendations=recommendations,
            risk_score=risk_score,
            risk_level=risk_level
        )
        return answer

    except Exception as exc:
        return f"Executive Consultant çıktısı üretilemedi: {exc}"


def create_executive_consultant_docx(company_name, mode, content):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    safe_company = str(company_name).replace(" ", "_").replace("/", "_")
    file_name = os.path.join(EXPORT_DIR, f"OptiFlow_{safe_company}_{mode.replace(' ', '_')}.docx")

    doc = Document()
    title = doc.add_heading("OptiFlow Executive Consultant", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f"{company_name} | {mode}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    for line in str(content).split("\n"):
        clean = line.strip()
        if not clean:
            doc.add_paragraph("")
            continue

        if clean[0:2].replace(".", "").isdigit() or clean.endswith(":"):
            doc.add_heading(clean, level=2)
        else:
            p = doc.add_paragraph(clean)
            for run in p.runs:
                run.font.size = Pt(10.5)

    doc.save(file_name)
    return file_name


def create_executive_consultant_txt(company_name, mode, content):
    safe_company = str(company_name).replace(" ", "_").replace("/", "_")
    file_name = os.path.join(EXPORT_DIR, f"OptiFlow_{safe_company}_{mode.replace(' ', '_')}.txt")

    with open(file_name, "w", encoding="utf-8") as file:
        file.write(str(content))

    return file_name


def render_ai_executive_consultant(
    company_name,
    sector,
    score,
    maturity,
    company_metrics,
    benchmark_result,
    financial_result,
    recommendations,
    risk_score,
    risk_level,
    active_plan,
    active_rules
):
    st.markdown("## V16 Executive Consultant")
    st.caption("CEO, COO, CFO ve yönetim kurulu seviyesinde profesyonel danışmanlık çıktıları üretir.")

    if not active_rules.get("ai", False):
        render_locked_feature("Executive Consultant", active_plan)
        return

    mode = st.selectbox(
        "Danışmanlık çıktısı seç",
        [
            "CEO Summary",
            "COO Operations Brief",
            "CFO Financial Brief",
            "90-Day Transformation Plan",
            "Board Presentation Notes",
            "Sales Proposal"
        ]
    )

    st.markdown("### Analiz Özeti")
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("OptiFlow Score", f"{score}/100")
    c2.metric("Risk", risk_level)
    c3.metric("Yıllık Tasarruf", money_fmt(financial_result.get("Tahmini Yıllık Tasarruf", 0)))
    c4.metric("ROI", pct_fmt(financial_result.get("ROI (%)", 0)))

    with st.expander("Analiz Verisi"):
        st.text(
            build_executive_consultant_context(
                company_name,
                sector,
                score,
                maturity,
                company_metrics,
                benchmark_result,
                financial_result,
                recommendations,
                risk_score,
                risk_level
            )
        )

    if "v16_consultant_output" not in st.session_state:
        st.session_state.v16_consultant_output = ""

    if st.button("Executive Consultant Çıktısı Üret", type="primary"):
        with st.spinner("OptiFlow Executive Consultant hazırlanıyor..."):
            st.session_state.v16_consultant_output = generate_executive_consultant_output(
                mode=mode,
                company_name=company_name,
                sector=sector,
                score=score,
                maturity=maturity,
                company_metrics=company_metrics,
                benchmark_result=benchmark_result,
                financial_result=financial_result,
                recommendations=recommendations,
                risk_score=risk_score,
                risk_level=risk_level
            )

    if st.session_state.v16_consultant_output:
        st.markdown("### Executive Consultant Output")
        st.markdown(
            f"""
            <div style="padding:24px;border-radius:18px;background:#f8fafc;border-left:6px solid #2563eb;">
            {st.session_state.v16_consultant_output.replace(chr(10), '<br>')}
            </div>
            """,
            unsafe_allow_html=True
        )

        docx_file = create_executive_consultant_docx(
            company_name=company_name,
            mode=mode,
            content=st.session_state.v16_consultant_output
        )

        txt_file = create_executive_consultant_txt(
            company_name=company_name,
            mode=mode,
            content=st.session_state.v16_consultant_output
        )

        col_download1, col_download2, col_mail = st.columns(3)

        with col_download1:
            with open(docx_file, "rb") as file:
                st.download_button(
                    "Executive Consultant DOCX İndir",
                    data=file,
                    file_name=os.path.basename(docx_file),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        with col_download2:
            with open(txt_file, "rb") as file:
                st.download_button(
                    "Executive Consultant TXT İndir",
                    data=file,
                    file_name=os.path.basename(txt_file),
                    mime="text/plain"
                )

        with col_mail:
            send_to = st.text_input("Mail alıcısı", value=user_email or "", key="v16_mail_to")
            if st.button("Executive Output Mail Gönder"):
                ok, msg = send_email_with_attachments(
                    to_email=send_to,
                    subject=f"OptiFlow Executive Consultant - {company_name}",
                    body=f"""Merhaba,

{company_name} için hazırlanan OptiFlow Executive Consultant çıktısı ekte yer almaktadır.

Çıktı türü: {mode}

Saygılarımızla,
OptiFlow Consulting""",
                    attachment_paths=[docx_file, txt_file]
                )

                save_email_log(
                    user_email=user_email,
                    to_email=send_to,
                    subject=f"OptiFlow Executive Consultant - {company_name}",
                    status="success" if ok else "failed",
                    message=msg,
                    company_name=company_name
                )

                if ok:
                    st.success("Executive Consultant çıktısı e-posta ile gönderildi.")
                else:
                    st.error(msg)



# ============================================================
# V14 REPORT DELIVERY CENTER + GMAIL SMTP
# ============================================================

def get_gmail_settings():
    sender = st.secrets.get("GMAIL_SENDER", "")
    password = st.secrets.get("GMAIL_APP_PASSWORD", "")
    return sender, password


def send_email_with_attachments(to_email, subject, body, attachment_paths=None):
    sender, password = get_gmail_settings()

    if not sender or not password:
        return False, "Gmail SMTP bilgileri eksik. Secrets içine GMAIL_SENDER ve GMAIL_APP_PASSWORD eklenmeli."

    if not to_email:
        return False, "Alıcı e-posta adresi bulunamadı."

    attachment_paths = attachment_paths or []

    try:
        msg = EmailMessage()
        msg["From"] = sender
        msg["To"] = to_email
        msg["Subject"] = subject
        msg.set_content(body)

        for path in attachment_paths:
            try:
                if path and os.path.exists(path):
                    with open(path, "rb") as file:
                        data = file.read()

                    file_name = os.path.basename(path)

                    if file_name.lower().endswith(".pdf"):
                        maintype, subtype = "application", "pdf"
                    elif file_name.lower().endswith(".xlsx"):
                        maintype, subtype = "application", "vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    elif file_name.lower().endswith(".pptx"):
                        maintype, subtype = "application", "vnd.openxmlformats-officedocument.presentationml.presentation"
                    else:
                        maintype, subtype = "application", "octet-stream"

                    msg.add_attachment(
                        data,
                        maintype=maintype,
                        subtype=subtype,
                        filename=file_name
                    )
            except Exception:
                pass

        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
            smtp.login(sender, password)
            smtp.send_message(msg)

        return True, "E-posta başarıyla gönderildi."

    except Exception as exc:
        return False, str(exc)


def create_email_delivery_log_table_if_needed():
    supabase = get_supabase_client()
    if not supabase:
        return

    # Not: Tablo SQL ile oluşturulmalı. Bu fonksiyon sadece uygulama tarafında sessiz geçer.
    return


def save_email_log(user_email, to_email, subject, status, message, company_name=None):
    supabase = get_supabase_client()

    if not supabase:
        return None

    try:
        payload = {
            "user_email": user_email,
            "to_email": to_email,
            "subject": subject,
            "status": status,
            "message": str(message)[:900],
            "company_name": company_name
        }

        result = supabase.table("email_logs").insert(payload).execute()

        if result.data:
            return result.data[0].get("id")

    except Exception:
        # email_logs tablosu yoksa uygulama bozulmasın.
        return None

    return None


def load_user_reports(user_email):
    supabase = get_supabase_client()

    if not supabase or not user_email:
        return []

    try:
        result = (
            supabase.table("reports")
            .select("*")
            .eq("user_email", user_email)
            .order("created_at", desc=True)
            .execute()
        )
        return result.data or []
    except Exception:
        return []


def load_email_logs(user_email=None, admin=False):
    supabase = get_supabase_client()

    if not supabase:
        return []

    try:
        query = supabase.table("email_logs").select("*").order("created_at", desc=True)
        if not admin and user_email:
            query = query.eq("user_email", user_email)
        result = query.execute()
        return result.data or []
    except Exception:
        return []


def render_my_reports_center(user_email):
    st.markdown("## My Reports")
    st.caption("Kullanıcıya ait oluşturulan rapor geçmişi")

    reports = load_user_reports(user_email)

    if not reports:
        st.info("Henüz rapor geçmişi bulunmuyor. Report Center veya Excel Upload üzerinden rapor oluştur.")
        return

    df = pd.DataFrame(reports)
    visible_cols = [c for c in ["report_type", "file_name", "file_path", "created_at"] if c in df.columns]
    st.dataframe(df[visible_cols], use_container_width=True, hide_index=True)

    st.markdown("### Raporu E-posta ile Gönder")

    selected = st.selectbox(
        "Gönderilecek raporu seç",
        df["file_name"].tolist()
    )

    selected_row = df[df["file_name"] == selected].iloc[0].to_dict()
    to_email = st.text_input("Alıcı e-posta", value=user_email or "")

    subject = st.text_input(
        "Konu",
        value=f"OptiFlow Raporunuz Hazır - {selected_row.get('file_name', '')}"
    )

    body = st.text_area(
        "E-posta mesajı",
        value=f"""Merhaba,

OptiFlow analiz raporunuz hazırdır.

Ekli dosyada rapor çıktısını bulabilirsiniz.

Saygılarımızla,
OptiFlow Consulting"""
    )

    if st.button("Seçili Raporu Mail Gönder", type="primary"):
        path = selected_row.get("file_path")
        ok, msg = send_email_with_attachments(
            to_email=to_email,
            subject=subject,
            body=body,
            attachment_paths=[path]
        )

        save_email_log(
            user_email=user_email,
            to_email=to_email,
            subject=subject,
            status="success" if ok else "failed",
            message=msg
        )

        if ok:
            st.success(msg)
        else:
            st.error(msg)


def render_report_delivery_center(user_email, active_plan, active_rules):
    st.markdown("## Report Delivery Center")
    st.caption("PDF, PPT ve Excel raporlarını e-posta ile gönderme merkezi")

    if active_plan == "demo":
        st.warning("Demo planda e-posta ile rapor gönderimi kapalıdır.")
        render_locked_feature("Report Delivery", active_plan)
        return

    reports = load_user_reports(user_email)

    if not reports:
        st.info("Gönderilecek rapor bulunamadı. Önce Report Center veya Excel Upload üzerinden rapor oluştur.")
        return

    df = pd.DataFrame(reports)

    col1, col2 = st.columns(2)

    with col1:
        report_options = df["file_name"].tolist()
        selected_files = st.multiselect("Gönderilecek raporlar", report_options, default=report_options[:1])

    with col2:
        to_email = st.text_input("Alıcı e-posta", value=user_email or "")

    subject = st.text_input(
        "Konu",
        value="OptiFlow Consulting Raporlarınız Hazır"
    )

    body = st.text_area(
        "E-posta mesajı",
        value=f"""Merhaba,

OptiFlow Consulting analiz çıktılarınız hazırlanmıştır.

Ekli dosyalarda PDF, PowerPoint veya Excel raporlarınızı bulabilirsiniz.

Bu rapor; operasyonel performans, finansal etki, benchmark ve yönetim önerileri içermektedir.

Saygılarımızla,
OptiFlow Consulting"""
    )

    if st.button("Seçili Raporları E-posta ile Gönder", type="primary"):
        selected_paths = []
        for file_name in selected_files:
            row = df[df["file_name"] == file_name]
            if not row.empty:
                selected_paths.append(row.iloc[0].get("file_path"))

        ok, msg = send_email_with_attachments(
            to_email=to_email,
            subject=subject,
            body=body,
            attachment_paths=selected_paths
        )

        save_email_log(
            user_email=user_email,
            to_email=to_email,
            subject=subject,
            status="success" if ok else "failed",
            message=msg
        )

        if ok:
            st.success(msg)
        else:
            st.error(msg)

    st.markdown("### Gönderim Geçmişi")
    logs = load_email_logs(user_email=user_email, admin=False)

    if logs:
        log_df = pd.DataFrame(logs)
        cols = [c for c in ["to_email", "subject", "status", "message", "created_at"] if c in log_df.columns]
        st.dataframe(log_df[cols], use_container_width=True, hide_index=True)
    else:
        st.info("Henüz e-posta gönderim kaydı yok.")


def render_admin_email_logs(user_email):
    if not is_admin_user(user_email):
        st.error("Bu sayfa sadece admin kullanıcılar içindir.")
        st.stop()

    st.markdown("## Admin Email Logs")
    st.caption("Tüm kullanıcıların e-posta gönderim kayıtları")

    logs = load_email_logs(admin=True)

    if not logs:
        st.info("Henüz e-posta log kaydı yok. email_logs tablosu yoksa SQL dosyasını çalıştır.")
        return

    df = pd.DataFrame(logs)
    cols = [c for c in ["user_email", "to_email", "company_name", "subject", "status", "message", "created_at"] if c in df.columns]

    c1, c2, c3 = st.columns(3)
    c1.metric("Toplam Mail", len(df))
    if "status" in df.columns:
        c2.metric("Başarılı", int((df["status"] == "success").sum()))
        c3.metric("Başarısız", int((df["status"] == "failed").sum()))

    st.dataframe(df[cols], use_container_width=True, hide_index=True)



# ============================================================
# V13 ADMIN ANALYTICS DASHBOARD
# ============================================================

def _df_or_empty(data):
    try:
        return pd.DataFrame(data or [])
    except Exception:
        return pd.DataFrame()


def render_admin_analytics(user_email):
    if not is_admin_user(user_email):
        st.error("Bu sayfa sadece admin kullanıcılar içindir.")
        st.stop()

    st.markdown("## Admin Analytics")
    st.caption("Platform performans, kullanım, rapor ve müşteri aktivitesi yönetim ekranı")

    users = admin_load_table("users")
    projects = admin_load_table("projects")
    reports = admin_load_table("reports")
    subscriptions = admin_load_table("subscriptions")

    users_df = _df_or_empty(users)
    projects_df = _df_or_empty(projects)
    reports_df = _df_or_empty(reports)
    subs_df = _df_or_empty(subscriptions)

    total_users = len(users_df)
    total_projects = len(projects_df)
    total_reports = len(reports_df)

    active_subs = 0
    if not subs_df.empty and "status" in subs_df.columns:
        active_subs = int(subs_df["status"].astype(str).str.lower().isin(["active", "trialing"]).sum())

    total_saving = 0
    if not projects_df.empty and "annual_saving" in projects_df.columns:
        total_saving = pd.to_numeric(projects_df["annual_saving"], errors="coerce").fillna(0).sum()

    avg_score = 0
    if not projects_df.empty and "score" in projects_df.columns:
        avg_score = pd.to_numeric(projects_df["score"], errors="coerce").fillna(0).mean()

    pdf_count = 0
    ppt_count = 0
    excel_count = 0
    if not reports_df.empty and "report_type" in reports_df.columns:
        report_types = reports_df["report_type"].astype(str).str.lower()
        pdf_count = int((report_types == "pdf").sum())
        ppt_count = int((report_types == "ppt").sum() + (report_types == "pptx").sum())
        excel_count = int((report_types == "excel").sum())

    k1, k2, k3, k4, k5 = st.columns(5)
    k1.metric("Toplam Kullanıcı", total_users)
    k2.metric("Aktif Abone", active_subs)
    k3.metric("Toplam Analiz", total_projects)
    k4.metric("Toplam Rapor", total_reports)
    k5.metric("Ortalama Score", f"{avg_score:.1f}")

    k6, k7, k8, k9 = st.columns(4)
    k6.metric("Toplam Tasarruf Potansiyeli", money_fmt(total_saving))
    k7.metric("PDF", pdf_count)
    k8.metric("PPT", ppt_count)
    k9.metric("Excel", excel_count)

    st.markdown("---")

    chart_col1, chart_col2 = st.columns(2)

    with chart_col1:
        st.markdown("### Plan Dağılımı")
        if not users_df.empty and "plan" in users_df.columns:
            plan_counts = (
                users_df["plan"]
                .fillna("demo")
                .astype(str)
                .str.lower()
                .value_counts()
                .reset_index()
            )
            plan_counts.columns = ["Plan", "Kullanıcı Sayısı"]
            fig = px.pie(
                plan_counts,
                names="Plan",
                values="Kullanıcı Sayısı",
                hole=0.45,
                title="Kullanıcı Plan Dağılımı"
            )
            fig.update_layout(height=360, margin=dict(l=20, r=20, t=55, b=20))
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("Plan verisi bulunamadı.")

    with chart_col2:
        st.markdown("### Rapor Tipleri")
        if not reports_df.empty and "report_type" in reports_df.columns:
            report_counts = (
                reports_df["report_type"]
                .fillna("unknown")
                .astype(str)
                .str.lower()
                .value_counts()
                .reset_index()
            )
            report_counts.columns = ["Rapor Tipi", "Adet"]
            fig = px.bar(
                report_counts,
                x="Rapor Tipi",
                y="Adet",
                text="Adet",
                title="Rapor Kullanımı"
            )
            fig.update_layout(height=360, margin=dict(l=20, r=20, t=55, b=20))
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("Rapor verisi bulunamadı.")

    st.markdown("### Son 30 Gün Aktivite")

    if not projects_df.empty and "created_at" in projects_df.columns:
        temp = projects_df.copy()
        temp["created_at"] = pd.to_datetime(temp["created_at"], errors="coerce")
        temp = temp.dropna(subset=["created_at"])

        if not temp.empty:
            last_30 = temp[temp["created_at"] >= (pd.Timestamp.utcnow() - pd.Timedelta(days=30))]
            daily = (
                last_30
                .assign(date=last_30["created_at"].dt.date)
                .groupby("date")
                .size()
                .reset_index(name="Analiz Sayısı")
            )

            if not daily.empty:
                fig = px.line(
                    daily,
                    x="date",
                    y="Analiz Sayısı",
                    markers=True,
                    title="Günlük Analiz Aktivitesi"
                )
                fig.update_layout(height=360, margin=dict(l=20, r=20, t=55, b=20))
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.info("Son 30 günde analiz verisi yok.")
        else:
            st.info("Aktivite tarih verisi okunamadı.")
    else:
        st.info("Aktivite verisi bulunamadı.")

    st.markdown("### Sektör Bazlı Kullanım")

    if not projects_df.empty and "sector" in projects_df.columns:
        sector_counts = (
            projects_df["sector"]
            .fillna("Unknown")
            .astype(str)
            .value_counts()
            .reset_index()
        )
        sector_counts.columns = ["Sektör", "Analiz Sayısı"]

        fig = px.bar(
            sector_counts,
            x="Sektör",
            y="Analiz Sayısı",
            text="Analiz Sayısı",
            title="En Çok Analiz Yapılan Sektörler"
        )
        fig.update_layout(height=380, margin=dict(l=20, r=20, t=55, b=20))
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("Sektör verisi bulunamadı.")

    st.markdown("### En Yüksek Tasarruf Potansiyeli Olan Projeler")

    if not projects_df.empty:
        display_cols = [c for c in ["company_name", "sector", "score", "risk_level", "annual_saving", "created_at", "user_email"] if c in projects_df.columns]
        top_projects = projects_df.copy()

        if "annual_saving" in top_projects.columns:
            top_projects["annual_saving"] = pd.to_numeric(top_projects["annual_saving"], errors="coerce").fillna(0)
            top_projects = top_projects.sort_values("annual_saving", ascending=False).head(10)

        st.dataframe(
            top_projects[display_cols],
            use_container_width=True,
            hide_index=True
        )
    else:
        st.info("Proje verisi bulunamadı.")

    st.markdown("### Kullanıcı Aktivite Tablosu")

    if not users_df.empty:
        user_activity = users_df.copy()

        if not projects_df.empty and "user_email" in projects_df.columns:
            project_counts = projects_df.groupby("user_email").size().reset_index(name="analysis_count")
            user_activity = user_activity.merge(
                project_counts,
                how="left",
                left_on="email",
                right_on="user_email"
            )
            user_activity["analysis_count"] = user_activity["analysis_count"].fillna(0).astype(int)
        else:
            user_activity["analysis_count"] = 0

        if not reports_df.empty and "user_email" in reports_df.columns:
            report_counts = reports_df.groupby("user_email").size().reset_index(name="report_count")
            user_activity = user_activity.merge(
                report_counts,
                how="left",
                left_on="email",
                right_on="user_email",
                suffixes=("", "_report")
            )
            user_activity["report_count"] = user_activity["report_count"].fillna(0).astype(int)
        else:
            user_activity["report_count"] = 0

        cols = [c for c in ["email", "full_name", "plan", "plan_status", "analysis_limit", "analysis_count", "report_count", "created_at"] if c in user_activity.columns]

        st.dataframe(
            user_activity[cols],
            use_container_width=True,
            hide_index=True
        )
    else:
        st.info("Kullanıcı verisi bulunamadı.")

    st.markdown("### Admin Export")

    if st.button("Admin Analytics Excel Oluştur", type="primary"):
        file_name = os.path.join(EXPORT_DIR, "OptiFlow_Admin_Analytics.xlsx")

        with pd.ExcelWriter(file_name, engine="openpyxl") as writer:
            users_df.to_excel(writer, sheet_name="Users", index=False)
            projects_df.to_excel(writer, sheet_name="Projects", index=False)
            reports_df.to_excel(writer, sheet_name="Reports", index=False)
            subs_df.to_excel(writer, sheet_name="Subscriptions", index=False)

        with open(file_name, "rb") as file:
            st.download_button(
                "Admin Analytics Excel İndir",
                data=file,
                file_name="OptiFlow_Admin_Analytics.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )


# ============================================================
# AUTH GATE
# ============================================================

if "demo_mode" not in st.session_state:
    st.session_state.demo_mode = False

if not is_logged_in() and not st.session_state.get("demo_mode", False):
    render_public_landing()
    st.stop()

user_email = current_user_email()
active_plan, active_rules, user_profile, subscription = resolve_active_plan(user_email)
can_analyze, analyses_used, analyses_limit = can_create_analysis(user_email, active_rules)

render_user_bar(user_email, active_plan, active_rules, analyses_used, analyses_limit)


if page == "Ana Sayfa":
    st.markdown(
        """
<div class="hero">
    <div class="hero-title">OptiFlow Consulting</div>
    <div class="hero-subtitle">
        Operational Excellence Intelligence Platform with Plotly Executive Dashboards, KPI Diagnostics and Enterprise Reporting.
    </div>
    <span class="hero-badge">Plotly Dashboard</span>
    <span class="hero-badge">KPI Diagnostics</span>
    <span class="hero-badge">Financial Impact</span>
    <span class="hero-badge">Consulting</span>
    <span class="hero-badge">PDF & PPT Export</span>
</div>
        """,
        unsafe_allow_html=True
    )

    col_a, col_b, col_c = st.columns(3)

    with col_a:
        st.info("1. Upload Excel or enter operational KPIs.")
    with col_b:
        st.info("2. Quantify financial impact and operational risk.")
    with col_c:
        st.info("3. Export executive PDF and PPT reports.")

    st.stop()


st.markdown(
    """
<div class="hero">
    <div class="hero-title">OptiFlow Consulting</div>
    <div class="hero-subtitle">
        Commercial Operations Excellence Platform | Plotly Dashboard | Benchmark Intelligence | PDF & PPT Export
    </div>
</div>
    """,
    unsafe_allow_html=True
)


with st.sidebar:
    st.markdown("---")
    st.subheader("Analiz Ayarları")

    company_name = st.text_input("Müşteri Firma", "Demo Firma A.Ş.")

    sector = st.selectbox(
        "Sektör",
        ["Tekstil", "Otomotiv", "Gida", "Lojistik", "Genel Uretim"]
    )

    input_mode = st.radio(
        "Veri Giriş Tipi",
        ["Manuel Giriş", "Excel Upload"]
    )

    uploaded_excel = None
    excel_result = None

    if input_mode == "Excel Upload":
        uploaded_excel = st.file_uploader("Müşteri Excel dosyası yükle", type=["xlsx", "xls"])

        if create_template_excel:
            template_path = create_template_excel()
            with open(template_path, "rb") as file:
                st.download_button(
                    "Excel Şablonu İndir",
                    data=file,
                    file_name="OptiFlow_Excel_Template.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

        if uploaded_excel and process_uploaded_excel:
            excel_result = process_uploaded_excel(uploaded_excel)
            st.success("Excel başarıyla okundu.")

    st.markdown("---")

    if excel_result:
        metrics = excel_result["metrics"]
        financial_inputs = excel_result["financial_inputs"]

        wait_rate = float(metrics["wait_rate"])
        oee = float(metrics["oee"])
        defect_rate = float(metrics["defect_rate"])
        line_balance_loss = float(metrics["line_balance_loss"])
        capacity_score = float(metrics["capacity_score"])

        total_wait_minutes = float(financial_inputs["total_wait_minutes"])
        hourly_labor_cost = float(financial_inputs["hourly_labor_cost"])
        working_days = int(financial_inputs["working_days"])
        improvement_rate = int(financial_inputs["improvement_rate"])

        st.info("KPI değerleri Excel verisine göre hesaplandı.")
    else:
        st.subheader("Operasyonel KPI")
        wait_rate = st.number_input("Bekleme Oranı (%)", min_value=0.0, max_value=100.0, value=27.0, step=1.0)
        oee = st.number_input("OEE (%)", min_value=0.0, max_value=100.0, value=68.0, step=1.0)
        defect_rate = st.number_input("Hata / Fire Oranı (%)", min_value=0.0, max_value=100.0, value=5.0, step=1.0)
        line_balance_loss = st.number_input("Hat Denge Kaybı (%)", min_value=0.0, max_value=100.0, value=22.0, step=1.0)
        capacity_score = st.number_input("Kapasite Skoru (%)", min_value=0.0, max_value=100.0, value=75.0, step=1.0)

        st.markdown("---")
        st.subheader("Finansal Varsayımlar")
        total_wait_minutes = st.number_input("Toplam Bekleme Süresi (dk/gün)", min_value=0.0, value=120.0, step=10.0)
        hourly_labor_cost = st.number_input("Saatlik İşçilik Maliyeti (TL)", min_value=0.0, value=250.0, step=25.0)
        working_days = st.number_input("Aylık Çalışma Günü", min_value=1, max_value=31, value=22)
        improvement_rate = st.slider("Bekleme İyileştirme Oranı (%)", 5, 60, 20)

    st.markdown("---")
    save_project = st.checkbox("Analizi müşteri arşivine kaydet", value=True)


company_metrics = {
    "wait_rate": wait_rate,
    "oee": oee,
    "defect_rate": defect_rate,
    "line_balance_loss": line_balance_loss
}

benchmark_result = benchmark_from_database(company_metrics, sector)

score = calculate_optiflow_score(
    efficiency_score=100 - wait_rate,
    oee_score=oee,
    capacity_score=capacity_score,
    flow_score=100 - line_balance_loss,
    quality_score=100 - defect_rate
)

maturity = get_maturity_comment(score)

financial_result = calculate_financial_impact(
    total_wait_minutes=total_wait_minutes,
    improvement_rate=improvement_rate,
    hourly_labor_cost=hourly_labor_cost,
    working_days_per_month=working_days
)

recommendations = generate_recommendations(
    wait_rate=wait_rate,
    oee=oee,
    line_balance_loss=line_balance_loss
)

risk_score, risk_level, risk_color, risk_bg = risk_status(
    wait_rate,
    oee,
    defect_rate,
    line_balance_loss
)


if page == "Dashboard":
    render_plotly_dashboard(
        company_name=company_name,
        score=score,
        maturity=maturity,
        company_metrics=company_metrics,
        benchmark_result=benchmark_result,
        financial_result=financial_result,
        risk_score=risk_score,
        risk_level=risk_level
    )


elif page == "Danışman Asistanı":
    if not active_rules.get("ai", False):
        render_locked_feature("Danışman Asistanı", active_plan)
    else:
        render_ai_copilot(
            company_name=company_name,
            sector=sector,
            score=score,
            maturity=maturity,
            company_metrics=company_metrics,
            financial_result=financial_result,
            recommendations=recommendations,
            risk_score=risk_score,
            risk_level=risk_level
        )




elif page == "Executive Consultant":
    render_ai_executive_consultant(
        company_name=company_name,
        sector=sector,
        score=score,
        maturity=maturity,
        company_metrics=company_metrics,
        benchmark_result=benchmark_result,
        financial_result=financial_result,
        recommendations=recommendations,
        risk_score=risk_score,
        risk_level=risk_level,
        active_plan=active_plan,
        active_rules=active_rules
    )


elif page == "Analysis":
    left, right = st.columns(2)

    with left:
        st.markdown("### Benchmark Karşılaştırması")
        st.dataframe(benchmark_result, use_container_width=True)

    with right:
        st.markdown("### Finansal Etki Analizi")
        st.dataframe(financial_result, use_container_width=True)

    if excel_result:
        st.markdown("### Yüklenen Excel Verisi")
        st.dataframe(excel_result["dataframe"], use_container_width=True)

        if excel_result.get("bottleneck"):
            st.markdown("### Otomatik Darboğaz Tespiti")
            st.info(
                f"Kritik darboğaz: {excel_result['bottleneck']['process']} - "
                f"{excel_result['bottleneck']['wait_minutes']} dk bekleme"
            )

        if excel_result.get("pareto_data") is not None:
            st.markdown("### Pareto Verisi")
            st.dataframe(excel_result["pareto_data"], use_container_width=True)

    st.markdown("### Yönetim İçin Öncelikli Aksiyonlar")
    for i, rec in enumerate(recommendations, start=1):
        st.markdown(f"**{i}.** {rec}")




elif page == "Excel Upload":
    render_v12_excel_upload_center(
        company_name=company_name,
        sector=sector,
        active_plan=active_plan,
        active_rules=active_rules
    )


elif page == "Benchmark Center":
    st.markdown("### Benchmark Center")

    db = load_benchmark_db()
    st.write("Sektör benchmark değerlerini buradan güncelleyebilirsin.")

    selected_sector = st.selectbox("Benchmark sektörü seç", list(db.keys()))

    c1, c2, c3, c4 = st.columns(4)

    with c1:
        new_wait = st.number_input("Sektör Bekleme Oranı", value=float(db[selected_sector].get("wait_rate", 0)))
    with c2:
        new_oee = st.number_input("Sektör OEE", value=float(db[selected_sector].get("oee", 0)))
    with c3:
        new_defect = st.number_input("Sektör Hata Oranı", value=float(db[selected_sector].get("defect_rate", 0)))
    with c4:
        new_balance = st.number_input("Sektör Hat Denge Kaybı", value=float(db[selected_sector].get("line_balance_loss", 0)))

    if st.button("Benchmark Verisini Güncelle"):
        db[selected_sector] = {
            "wait_rate": new_wait,
            "oee": new_oee,
            "defect_rate": new_defect,
            "line_balance_loss": new_balance
        }
        save_benchmark_db(db)
        st.success("Benchmark verisi güncellendi.")

    st.markdown("### Benchmark Veritabanı")
    st.json(db)


elif page == "Clients":
    st.markdown("### My Client Projects")

    projects = load_user_projects(user_email)

    if not projects:
        st.info("No saved projects yet. Create a report from Report Center to save a project.")
    else:
        st.dataframe(
            pd.DataFrame(projects)[[
                "company_name",
                "sector",
                "score",
                "risk_level",
                "annual_saving",
                "created_at"
            ]],
            use_container_width=True,
            hide_index=True
        )

        selected_company = st.selectbox(
            "Select company",
            sorted(list(set([p.get("company_name", "-") for p in projects])))
        )

        for project in [p for p in projects if p.get("company_name") == selected_company]:
            with st.expander(f"{project.get('created_at')} | Score: {project.get('score')}/100"):
                st.json(project)


elif page == "Report Center":
    st.markdown("### Enterprise Report Center")
    st.write("PDF, PowerPoint ve Excel çıktıları plan seviyesine göre açılır.")

    send_mail_after_report = st.checkbox("Rapor oluşturulunca e-posta gönder", value=False)
    report_mail_to = st.text_input("Rapor alıcı e-posta", value=user_email or "")

    if not can_analyze:
        st.error("Analysis limit reached for your current plan.")
        render_locked_feature("Additional analyses", active_plan)
        st.stop()

    col_pdf, col_ppt = st.columns(2)

    with col_pdf:
        if not active_rules.get("pdf", False):
            render_locked_feature("PDF / Excel Export", active_plan)
        elif st.button("Enterprise PDF Raporu Oluştur", type="primary"):
            with st.spinner("OptiFlow Consulting PDF raporu hazırlanıyor..."):
                consulting_report = generate_consulting_report(
                    sector=sector,
                    company_metrics=company_metrics,
                    benchmark_result=benchmark_result,
                    financial_result=financial_result,
                    maturity=maturity,
                    recommendations=recommendations
                )

                pdf_file = create_enterprise_pdf(
                    company_name=company_name,
                    sector=sector,
                    score=score,
                    maturity=maturity,
                    company_metrics=company_metrics,
                    benchmark_result=benchmark_result,
                    financial_result=financial_result,
                    recommendations=recommendations,
                    consulting_report=consulting_report
                )

                excel_file = create_excel_report(
                    company_name=company_name,
                    sector=sector,
                    score=score,
                    maturity=maturity,
                    company_metrics=company_metrics,
                    benchmark_result=benchmark_result,
                    financial_result=financial_result,
                    recommendations=recommendations,
                    risk_score=risk_score,
                    risk_level=risk_level
                )

                supabase_project_id = save_project_to_supabase(
                    user_email=user_email,
                    company_name=company_name,
                    sector=sector,
                    score=score,
                    risk_level=risk_level,
                    financial_result=financial_result
                )

                if supabase_project_id:
                    save_report_to_supabase(
                        user_email=user_email,
                        project_id=supabase_project_id,
                        report_type="pdf",
                        file_name=os.path.basename(pdf_file),
                        file_path=pdf_file
                    )
                    save_report_to_supabase(
                        user_email=user_email,
                        project_id=supabase_project_id,
                        report_type="excel",
                        file_name=os.path.basename(excel_file),
                        file_path=excel_file
                    )

                project_path = None
                if save_project:
                    project_path = save_project_record(
                        company_name=company_name,
                        sector=sector,
                        score=score,
                        maturity=maturity,
                        company_metrics=company_metrics,
                        financial_result=financial_result,
                        recommendations=recommendations
                    )

            if send_mail_after_report:
                ok, mail_msg = send_email_with_attachments(
                    to_email=report_mail_to,
                    subject=f"OptiFlow Consulting Raporunuz Hazır - {company_name}",
                    body=f"""Merhaba,

{company_name} için OptiFlow Consulting analiz raporunuz hazırlanmıştır.

Ekli dosyalarda PDF ve Excel çıktılarınızı bulabilirsiniz.

Saygılarımızla,
OptiFlow Consulting""",
                    attachment_paths=[pdf_file, excel_file]
                )

                save_email_log(
                    user_email=user_email,
                    to_email=report_mail_to,
                    subject=f"OptiFlow Consulting Raporunuz Hazır - {company_name}",
                    status="success" if ok else "failed",
                    message=mail_msg,
                    company_name=company_name
                )

                if ok:
                    st.success("Rapor oluşturuldu ve e-posta gönderildi.")
                else:
                    st.warning(f"Rapor oluşturuldu fakat e-posta gönderilemedi: {mail_msg}")
            else:
                st.success("Enterprise PDF raporu başarıyla oluşturuldu.")

            if project_path:
                st.info(f"Müşteri analizi arşive kaydedildi: {project_path}")

            with open(pdf_file, "rb") as file:
                st.download_button(
                    label="Enterprise PDF İndir",
                    data=file,
                    file_name=f"OptiFlow_{company_name.replace(' ', '_')}_Enterprise_Report.pdf",
                    mime="application/pdf"
                )

            if "excel_file" in locals() and excel_file:
                with open(excel_file, "rb") as file:
                    st.download_button(
                        label="Enterprise Excel Data İndir",
                        data=file,
                        file_name=os.path.basename(excel_file),
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )

    with col_ppt:
        if not active_rules.get("ppt", False):
            render_locked_feature("PowerPoint Export", active_plan)
        elif st.button("Executive PPTX Oluştur"):
            if create_enterprise_pptx is None:
                st.error("PPTX modülü bulunamadı. modules/ppt_engine.py dosyasını ekle.")
            else:
                with st.spinner("PowerPoint raporu hazırlanıyor..."):
                    ppt_file = os.path.join(
                        EXPORT_DIR,
                        f"OptiFlow_{company_name.replace(' ', '_')}_Executive_Deck.pptx"
                    )

                    ppt_file = create_enterprise_pptx(
                        company_name=company_name,
                        sector=sector,
                        score=score,
                        maturity=maturity,
                        company_metrics=company_metrics,
                        financial_result=financial_result,
                        recommendations=recommendations,
                        output_file=ppt_file
                    )

                st.success("PowerPoint raporu oluşturuldu.")

                with open(ppt_file, "rb") as file:
                    st.download_button(
                        label="Executive PPTX İndir",
                        data=file,
                        file_name=os.path.basename(ppt_file),
                        mime="application/vnd.openxmlformats-officedocument.presentationml.presentation"
                    )


elif page == "Billing":
    st.markdown("### Billing & Subscription")

    st.write(f"Current user: **{user_email}**")
    st.write(f"Current plan: **{active_rules.get('label', active_plan)}**")
    st.write(f"Subscription status: **{subscription.get('status', 'inactive')}**")
    st.write(f"Usage: **{analyses_used}/{analyses_limit if analyses_limit < 999999 else 'Unlimited'} analyses**")

    st.markdown("### Upgrade Plans")

    c1, c2, c3 = st.columns(3)

    with c1:
        st.markdown("#### Starter")
        st.write("$49/month")
        st.write("10 analyses/month")
        st.write("PDF + Excel exports")
        link = st.secrets.get("STRIPE_STARTER_LINK", "")
        if link:
            st.link_button("Buy Starter", link)
        else:
            st.info("Stripe link not configured yet.")

    with c2:
        st.markdown("#### Professional")
        st.write("$149/month")
        st.write("100 analyses/month")
        st.write("PDF + PPT + Excel +")
        link = st.secrets.get("STRIPE_PROFESSIONAL_LINK", "")
        if link:
            st.link_button("Buy Professional", link)
        else:
            st.info("Stripe link not configured yet.")

    with c3:
        st.markdown("#### Enterprise")
        st.write("Custom")
        st.write("Unlimited usage")
        st.write("Team support + custom deployment")
        link = st.secrets.get("ENTERPRISE_CONTACT_LINK", "")
        if link:
            st.link_button("Contact Sales", link)
        else:
            st.info("Contact link not configured yet.")

    st.caption("After payment, subscription activation can be automated with Stripe webhooks. For now, update the subscriptions table manually or via webhook service.")


elif page == "Admin Panel":
    render_admin_dashboard(user_email)


elif page == "Admin Analytics":
    render_admin_analytics(user_email)


elif page == "My Reports":
    render_my_reports_center(user_email)


elif page == "Report Delivery":
    render_report_delivery_center(
        user_email=user_email,
        active_plan=active_plan,
        active_rules=active_rules
    )


elif page == "Email Logs":
    render_admin_email_logs(user_email)


elif page == "Client Portal":
    render_client_portal(user_email)


elif page == "Benchmark Intelligence":
    render_benchmark_intelligence(
        company_name=company_name,
        sector=sector,
        score=score,
        company_metrics=company_metrics,
        benchmark_result=benchmark_result,
        financial_result=financial_result,
        risk_level=risk_level
    )


elif page == "Proposal Generator":
    render_proposal_generator(
        company_name=company_name,
        sector=sector,
        score=score,
        maturity=maturity,
        company_metrics=company_metrics,
        financial_result=financial_result,
        recommendations=recommendations,
        risk_level=risk_level,
        active_plan=active_plan,
        active_rules=active_rules
    )
